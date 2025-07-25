import os
import pdfplumber  # PyMuPDF 대신 pdfplumber 사용
from docx import Document
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
import gc
import time
import re
from typing import List
from dotenv import load_dotenv

# .env 파일 로드
load_dotenv()

# === 경로 설정 ===
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCS_DIR = os.path.join(BASE_DIR, "docs")
DB_DIR = os.path.join(BASE_DIR, "db")

# 인덱스 분류별 경로 설정
GAS_INDEX_DIR = os.path.join(DB_DIR, "gas_index")
POWER_INDEX_DIR = os.path.join(DB_DIR, "power_index")
OTHER_INDEX_DIR = os.path.join(DB_DIR, "other_index")

# 필요한 폴더 생성
os.makedirs(DB_DIR, exist_ok=True)
os.makedirs(GAS_INDEX_DIR, exist_ok=True)
os.makedirs(POWER_INDEX_DIR, exist_ok=True)
os.makedirs(OTHER_INDEX_DIR, exist_ok=True)

# === OpenAI 임베딩 모델 초기화 ===
def create_embedding_model():
    """동적으로 배치 크기를 조정하여 임베딩 모델을 생성합니다."""
    # 초기 배치 크기
    chunk_size = 1000
    
    try:
        embedding_model = OpenAIEmbeddings(
            model="text-embedding-3-small",
            chunk_size=chunk_size,
            max_retries=3
        )
        return embedding_model
    except Exception as e:
        if "max_tokens_per_request" in str(e):
            # 토큰 제한 오류 시 배치 크기 줄임
            chunk_size = 500
            print(f"토큰 제한으로 인해 배치 크기를 {chunk_size}로 조정합니다.")
            return OpenAIEmbeddings(
                model="text-embedding-3-small",
                chunk_size=chunk_size,
                max_retries=3
            )
        else:
            raise e

embedding_model = create_embedding_model()

# === 텍스트 분할기 초기화 ===
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=400,
    chunk_overlap=80,
    length_function=len,
    separators=["\n\n", "\n", " ", ""]
)

# === 문서 읽기 ===
def read_docx(path):
    """DOCX 문서를 읽어서 텍스트를 추출합니다."""
    doc = Document(path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    text = latex_to_text(text)  # LaTeX 변환 추가
    print(f"[DOCX] {path} 추출 텍스트 예시:", text[:100])
    return text

def latex_to_text(text):
    """
    LaTeX 수식을 사람이 읽는 텍스트 수식으로 변환
    예: \frac{a}{b} → (a) / (b)
    """
    # \frac 변환 함수
    def frac_repl(match):
        return f"({match.group(1)}) / ({match.group(2)})"

    # LaTeX 블록(\[...\], $$...$$, $...$)을 찾아서 변환
    def latex_block_repl(match):
        latex_expr = match.group(1)
        # \left와 \right 제거 (괄호 크기 조정 명령어) - 더 포괄적으로 처리
        latex_expr = re.sub(r'\\left\s*\(', '(', latex_expr)
        latex_expr = re.sub(r'\\right\s*\)', ')', latex_expr)
        latex_expr = re.sub(r'\\left\s*\[', '[', latex_expr)
        latex_expr = re.sub(r'\\right\s*\]', ']', latex_expr)
        latex_expr = re.sub(r'\\left\s*\\{', '{', latex_expr)
        latex_expr = re.sub(r'\\right\s*\\}', '}', latex_expr)
        # \frac 변환
        latex_expr = re.sub(r'\\frac\{(.+?)\}\{(.+?)\}', frac_repl, latex_expr)
        # \times 변환
        latex_expr = latex_expr.replace(r'\times', '×')
        # 중괄호 제거
        latex_expr = latex_expr.replace('{', '').replace('}', '')
        return latex_expr

    # \[ ... \] 블록 변환
    text = re.sub(r'\\\[(.*?)\\\]', lambda m: latex_block_repl(m), text, flags=re.DOTALL)
    # $$ ... $$ 블록 변환
    text = re.sub(r'\$\$(.*?)\$\$', lambda m: latex_block_repl(m), text, flags=re.DOTALL)
    # $ ... $ 블록 변환
    text = re.sub(r'\$(.*?)\$', lambda m: latex_block_repl(m), text, flags=re.DOTALL)

    # 인라인 변환 (혹시 남아있을 경우)
    # \left와 \right 제거 - 더 포괄적으로 처리
    text = re.sub(r'\\left\s*\(', '(', text)
    text = re.sub(r'\\right\s*\)', ')', text)
    text = re.sub(r'\\left\s*\[', '[', text)
    text = re.sub(r'\\right\s*\]', ']', text)
    text = re.sub(r'\\left\s*\\{', '{', text)
    text = re.sub(r'\\right\s*\\}', '}', text)
    # \frac 변환
    text = re.sub(r'\\frac\{(.+?)\}\{(.+?)\}', frac_repl, text)
    text = text.replace(r'\times', '×')
    text = text.replace('{', '').replace('}', '')

    return text

def read_pdf(path):
    """PDF 문서를 읽어서 전체 텍스트를 추출합니다."""
    full_text = ""
    with pdfplumber.open(path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            try:
                page_text = page.extract_text()
                if page_text:
                    page_text = latex_to_text(page_text)
                    page_text = normalize_text(page_text)
                    full_text += page_text + "\n"
            except Exception as e:
                print(f"페이지 {page_num} 처리 중 오류: {str(e)}")
    if full_text:
        print(f"[PDF] {path} 텍스트 예시:", full_text[:100])
    return full_text

# === 텍스트 정규화 ===
def normalize_text(text):
    """텍스트 정규화 (한글-영문 공백, 수식 기호 보존)"""
    # 한글-영문 공백 정리
    text = re.sub(r'([가-힣])([A-Za-z0-9])', r'\1 \2', text)
    text = re.sub(r'([A-Za-z0-9])([가-힣])', r'\1 \2', text)

    # 수식 기호 보존: √ ± ≈ ∞ × ÷ π ² ³ ^ / = % 등
    math_symbols = "√±≈∞×÷π²³^=/%"

    # 허용 문자 정의: 문자, 숫자, 공백, 일부 수식 기호, 일반 문장부호
    allowed_chars = r'\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}' + re.escape(math_symbols)
    text = re.sub(rf'[^\w\s{re.escape(".,!?;:-()[]{}")}{"".join(math_symbols)}]', '', text)

    # 연속 공백 정리
    text = re.sub(r'\s+', ' ', text)

    return text.strip()

# === 파일 분류 함수 ===
def classify_file(file_name):
    """파일명에 따라 분류를 결정합니다."""
    if '도시가스' in file_name:
        return 'gas'
    elif '전력' in file_name:
        return 'power'
    else:
        return 'other'

def get_index_dir(category):
    """분류에 따른 인덱스 디렉토리를 반환합니다."""
    if category == 'gas':
        return GAS_INDEX_DIR
    elif category == 'power':
        return POWER_INDEX_DIR
    else:
        return OTHER_INDEX_DIR

def extract_metadata_from_filename(filename):
    base_name = os.path.splitext(filename)[0]
    parts = re.split(r'[_\-\s]', base_name)

    version = next((p for p in parts if re.match(r'20\d{2}', p)), None)
    region_list = ['서울', '부산', '대구', '광주', '인천', '대전', '울산','경기도', '강원도', '충청북도', '충청남도' ,'전라남도','전북특별자치도', '경상남도', '경상북도']
    region = next((p for p in parts if p in region_list), None)
    organization_map = {'도시가스': '도시가스', '전력': '전력'}
    organization = next((p for p in organization_map if p in parts), "기타")

    return {
        "version": version,
        "region": region,
        "organization": organization,
        "title": base_name
    }

# === 문서 처리 및 청킹 ===
def process_document(file_path):
    """문서를 읽고 LangChain 텍스트 분할기를 사용하여 청킹합니다."""
    filename = os.path.basename(file_path)
    
    # 문서 읽기
    if file_path.endswith(".pdf"):
        text = read_pdf(file_path)
    elif file_path.endswith(".docx"):
        text = read_docx(file_path)
    else:
        return []
    
    # 텍스트가 비어있으면 처리 중단
    if not text:
        print(f"⚠️  {filename}에서 텍스트를 추출하지 못했습니다. 건너뜁니다.")
        return []
        
    # 텍스트 정규화
    text = normalize_text(text)
    
    # LangChain 텍스트 분할기 사용
    chunks = text_splitter.split_text(text)
    
    # 파일명에서 추가 메타데이터 추출
    additional_metadata = extract_metadata_from_filename(filename)
    
    # 메타데이터가 포함된 LangChain 문서로 변환
    documents = []
    for i, chunk in enumerate(chunks):
        if len(chunk.strip()) > 50:  # 최소 길이 필터 (50자 이상)
            # 기본 메타데이터와 추가 메타데이터를 결합
            metadata = {
                "source": filename,
                "chunk_id": i,
                "file_path": file_path,
                "version": additional_metadata["version"],
                "region": additional_metadata["region"],
                "organization": additional_metadata["organization"],
                "title": additional_metadata["title"]
            }
            doc = LCDocument(
                page_content=chunk,
                metadata=metadata
            )
            documents.append(doc)
    
    print(f"[CHUNKS] {filename} → {len(documents)}개 생성")
    return documents

# === 벡터 인덱스 구축 ===
def build_langchain_vector_index():
    """문서들을 읽어서 분류별로 벡터 인덱스를 구축합니다."""
    print("📂 문서 인덱싱 시작")

    # 문서 목록 불러오기
    files = [f for f in os.listdir(DOCS_DIR) if f.endswith((".pdf", ".docx"))]
    if not files:
        print("❌ 인덱싱할 문서 없음")
        return

    print(f"총 {len(files)}개 문서 발견")

    # 분류별로 문서 그룹화
    gas_documents = []
    power_documents = []
    other_documents = []
    
    total_files = len(files)

    # 문서별로 처리 및 분류
    for i, file_name in enumerate(files, 1):
        file_path = os.path.join(DOCS_DIR, file_name)
        print(f"[{i}/{total_files}] 처리 중: {file_name}")
        
        # 파일 분류
        category = classify_file(file_name)
        print(f"  → 분류: {category}")
        
        documents = process_document(file_path)
        
        # 분류별로 문서 추가
        if category == 'gas':
            gas_documents.extend(documents)
        elif category == 'power':
            power_documents.extend(documents)
        else:
            other_documents.extend(documents)
        
        # 메모리 최적화
        del documents
        gc.collect()
        time.sleep(0.1)  # 시스템 여유 주기

    print(f"\n📊 전체 문서 수: {len(files)}개")
    print(f"🔖 Gas 문서: {len(gas_documents)}개")
    print(f"🔖 Power 문서: {len(power_documents)}개")
    print(f"🔖 Other 문서: {len(other_documents)}개")

    # 분류별로 인덱스 생성
    categories = [
        ('gas', gas_documents, 'Gas'),
        ('power', power_documents, 'Power'),
        ('other', other_documents, 'Other')
    ]
    
    for category, documents, category_name in categories:
        if len(documents) == 0:
            print(f"⚠️ {category_name} 문서가 없어 인덱스를 생성하지 않습니다.")
            continue
            
        print(f"\n🔧 {category_name} 인덱스 생성 중... (문서 수: {len(documents)}개)")
        
        # 배치 단위로 임베딩 처리
        try:
            vectorstore = FAISS.from_documents(documents, embedding_model)
        except Exception as e:
            if "max_tokens_per_request" in str(e):
                print(f"토큰 제한으로 인해 배치 크기를 500으로 조정합니다...")
                # 두 번째 시도: 500
                try:
                    medium_embedding_model = OpenAIEmbeddings(
                        model="text-embedding-3-small",
                        chunk_size=500,
                        max_retries=3
                    )
                    vectorstore = FAISS.from_documents(documents, medium_embedding_model)
                except Exception as e2:
                    if "max_tokens_per_request" in str(e2):
                        print(f"토큰 제한으로 인해 배치 크기를 100으로 조정합니다...")
                        # 최후 수단: 100
                        small_embedding_model = OpenAIEmbeddings(
                            model="text-embedding-3-small",
                            chunk_size=100,
                            max_retries=3
                        )
                        vectorstore = FAISS.from_documents(documents, small_embedding_model)
                    else:
                        raise e2
            else:
                raise e

        # 인덱스 저장
        index_dir = get_index_dir(category)
        try:
            vectorstore.save_local(index_dir)
            print(f"💾 {category_name} 인덱스 저장 완료: {index_dir}")
        except Exception as e:
            print(f"❌ {category_name} 인덱스 저장 오류: {e}")

    print("\n🎉 모든 분류별 인덱스 생성 완료")

# === 검색 시스템 ===
def search_query(query, top_k=10, category=None):
    """벡터 인덱스에서 유사한 문서를 검색합니다."""
    print(f"\n🔍 검색어: {query}")
    
    # 검색할 인덱스 결정
    if category is None:
        # 모든 인덱스에서 검색
        search_categories = ['gas', 'power', 'other']
        all_results = []
        
        for cat in search_categories:
            index_dir = get_index_dir(cat)
            index_faiss_path = os.path.join(index_dir, "index.faiss")
            
            if os.path.exists(index_faiss_path):
                try:
                    print(f"📂 {cat} 인덱스에서 검색 중...")
                    vectorstore = FAISS.load_local(index_dir, embedding_model, allow_dangerous_deserialization=True)
                    results = vectorstore.similarity_search_with_score(query, k=5)
                    all_results.extend(results)
                    print(f"  → {cat} 인덱스에서 {len(results)}개 결과 발견")
                except Exception as e:
                    print(f"⚠️ {cat} 인덱스 로드 실패: {e}")
            else:
                print(f"ℹ️ {cat} 인덱스가 존재하지 않습니다. (건너뜀)")
        
        # 모든 결과를 점수로 정렬
        if all_results:
            all_results = sorted(all_results, key=lambda x: x[1])
            results = all_results[:top_k]
        else:
            print("❌ 검색 결과가 없습니다.")
            return
        
    else:
        # 특정 카테고리에서만 검색
        index_dir = get_index_dir(category)
        index_faiss_path = os.path.join(index_dir, "index.faiss")
        
        if not os.path.exists(index_faiss_path):
            print(f"❌ {category} 인덱스를 찾을 수 없습니다: {index_faiss_path}")
            return
            
        try:
            vectorstore = FAISS.load_local(index_dir, embedding_model, allow_dangerous_deserialization=True)
            results = vectorstore.similarity_search_with_score(query, k=top_k)
            results = sorted(results, key=lambda x: x[1])  # 낮은 점수일수록 유사도 높음
        except Exception as e:
            print(f"❌ {category} 인덱스 로드 실패: {e}")
            return

    print(f"✅ 검색 완료 (총 {len(results)}개 결과)")

    print("\n📌 검색 결과:")
    for i, (doc, score) in enumerate(results, start=1):
        print(f"[{i}] 점수: {score:.4f}")
        # 검색어 주변 컨텍스트를 포함하여 출력
        content = doc.page_content
        if query in content:
            idx = content.find(query)
            start = max(0, idx - 150)  # 컨텍스트 범위 확장
            end = min(len(content), idx + 150)  # 컨텍스트 범위 확장
            print(f"...{content[start:end]}...")
        else:
            # 검색어가 없는 경우, 문장 단위로 출력
            sentences = content.split('.')
            print('. '.join(sentences[:3]) + '.')  # 처음 3개 문장만 출력
        print(f"     ⤷ 출처: {doc.metadata.get('source')} / 청크 ID: {doc.metadata.get('chunk_id')}")
        print("-" * 60)

# === 메인 실행 ===
if __name__ == "__main__":
    build_langchain_vector_index()
    
    # 인덱스 생성 완료 후 예시 검색 수행
    print("\n" + "="*60)
    print("🔍 검색 테스트")
    print("="*60)
    
    # 전체 검색
    search_query("기준열량")
    
    # 분류별 검색 예시
    #search_query("도시가스 공급 규정에서 도시가스회사의 의무는 무엇인가요?", category='gas')
    #search_query("가스요금은 어떻게 계산되나요?", category='gas')
    #search_query("전력 공급 규정", category='power')
    #search_query("기타 문서 검색", category='other')
    