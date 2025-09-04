# chat_ui_docinsight.py (DocInsight AI 스타일 적용 전체 버전)
import streamlit as st
import requests
import time
import re
import os
from dotenv import load_dotenv

# 문서 업로드 관련 추가 import
import json
import pandas as pd
from typing import List, Dict, Any
import tempfile
import shutil
from pathlib import Path
import unicodedata

# 통합 로그 관리 클래스
class UnifiedLogger:
    """한 영역에서 로그를 관리하는 클래스"""
    def __init__(self):
        self.main_container = None
        self.detail_container = None
        self.is_active = False
    
    def start(self, title="진행 상황"):
        """로깅 시작"""
        if not self.is_active:
            self.main_container = st.empty()
            self.detail_container = st.empty()
            self.is_active = True
        self.update_main(title, "시작", "🚀")
    
    def update_main(self, message, status="진행중", icon="📝"):
        """메인 상태 업데이트"""
        if self.main_container:
            status_color = {
                "시작": "#3b82f6",
                "진행중": "#f59e0b", 
                "완료": "#10b981",
                "실패": "#ef4444"
            }.get(status, "#6b7280")
            
            self.main_container.markdown(f"""
                <div style="background-color: rgba(59, 130, 246, 0.1); padding: 1rem; border-radius: 8px; border-left: 4px solid {status_color}; margin-bottom: 1rem;">
                    <div style="display: flex; align-items: center; font-weight: bold; color: {status_color};">
                        <span style="margin-right: 8px; font-size: 1.2rem;">{icon}</span>
                        <span>{message}</span>
                        <span style="margin-left: auto; font-size: 0.9rem;">({status})</span>
                    </div>
                </div>
            """, unsafe_allow_html=True)
    
    def update_detail(self, message, type="info"):
        """세부 진행상황 업데이트"""
        if self.detail_container:
            colors = {
                "info": "#3b82f6",
                "success": "#10b981", 
                "warning": "#f59e0b",
                "error": "#ef4444"
            }
            icons = {
                "info": "ℹ️",
                "success": "✅",
                "warning": "⚠️", 
                "error": "❌"
            }
            
            color = colors.get(type, "#6b7280")
            icon = icons.get(type, "📝")
            
            self.detail_container.markdown(f"""
                <div style="background-color: rgba(107, 114, 128, 0.05); padding: 0.75rem; border-radius: 6px; border-left: 3px solid {color}; margin-bottom: 0.5rem;">
                    <div style="display: flex; align-items: center; color: {color}; font-size: 0.9rem;">
                        <span style="margin-right: 6px;">{icon}</span>
                        <span>{message}</span>
                    </div>
                </div>
            """, unsafe_allow_html=True)
    
    def success(self, message):
        """성공 메시지"""
        self.update_main(message, "완료", "🎉")
        if self.detail_container:
            self.detail_container.empty()
    
    def error(self, message):
        """에러 메시지"""
        self.update_main(message, "실패", "❌")
        if self.detail_container:
            self.detail_container.empty()
    
    def finish(self, success_message=None, error_message=None):
        """로깅 종료"""
        if success_message:
            self.success(success_message)
        elif error_message:
            self.error(error_message)
        
        if self.detail_container:
            self.detail_container.empty()
        self.is_active = False
    
    def clear(self):
        """모든 메시지 정리"""
        if self.main_container:
            self.main_container.empty()
        if self.detail_container:
            self.detail_container.empty()
        self.is_active = False

# 문서 처리 및 벡터 인덱스 생성을 위한 모듈들
import pdfplumber
from docx import Document
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain_core.documents import Document as LCDocument
from langchain_experimental.text_splitter import SemanticChunker
import tiktoken  # 실제 토큰 카운트를 위한 라이브러리
import gc

# === UI 전용 설정 (통합 관리) ===
# OpenAI 모델 설정
OPENAI_EMBEDDING_MODEL = "text-embedding-3-small"  # 임베딩 모델명
OPENAI_MAX_RETRIES = 3  # API 재시도 횟수

# 임베딩 배치 크기 (API 호출 최적화) - 안정성과 효율성의 균형
EMBEDDING_BATCH_SIZE_INITIAL = 500     # 1차 시도 (적정 처리)
EMBEDDING_BATCH_SIZE_RETRY = 200       # 2차 시도 (안전)
EMBEDDING_BATCH_SIZE_FINAL = 100       # 3차 시도 (매우 안전)

# 페이지 설정
st.set_page_config(page_title="DocInsight AI", page_icon="📄", layout="wide")
load_dotenv()
API_URL = "http://localhost:8000/chat"
RELOAD_API_URL = "http://localhost:8000/reload-indexes"
DOCUMENT_API_URL = "http://localhost:8000/get-document"

# === 문서 업로드를 위한 함수들 (chat_ui.py에서 복사) ===

# 경로 설정
BASE_DIR = Path(__file__).parent.parent / "vectordb"
DOCS_DIR = BASE_DIR / "docs"
DB_DIR = BASE_DIR / "db"

# 인덱스 분류별 경로 설정
GAS_INDEX_DIR = DB_DIR / "gas_index"
POWER_INDEX_DIR = DB_DIR / "power_index"
OTHER_INDEX_DIR = DB_DIR / "other_index"

# 필요한 폴더 생성
DB_DIR.mkdir(exist_ok=True)
GAS_INDEX_DIR.mkdir(exist_ok=True)
POWER_INDEX_DIR.mkdir(exist_ok=True)
OTHER_INDEX_DIR.mkdir(exist_ok=True)

# 토큰 분할 처리를 위한 함수들
def get_token_encoding():
    """OpenAI 임베딩 모델에 맞는 토큰 인코딩을 반환합니다."""
    try:
        # text-embedding-3-small 모델에 맞는 인코딩 (cl100k_base)
        return tiktoken.get_encoding("cl100k_base")
    except Exception as e:
        # 로그 출력 제거 - 이미 상위 함수에서 처리됨
        return None

def count_actual_tokens(text: str) -> int:
    """실제 토큰 수를 정확히 계산합니다."""
    encoding = get_token_encoding()
    if encoding:
        try:
            return len(encoding.encode(text))
        except Exception as e:
            # 로그 출력 제거 - 이미 상위 함수에서 처리됨
            # 폴백: 추정 방식
            return int(len(text) * 0.4)
    else:
        # 폴백: 추정 방식
        return int(len(text) * 0.4)

def calculate_total_tokens(documents: list) -> int:
    """문서 리스트의 실제 총 토큰 수를 계산합니다."""
    total_text = "".join([doc.page_content for doc in documents])
    return count_actual_tokens(total_text)

def create_vectorstore_with_token_limit(documents, embedding_model, max_tokens=280000, logger=None):  # 28만 토큰으로 증가 (OpenAI 제한의 93%)
    """실제 토큰 수 기준으로 벡터스토어를 생성합니다."""
    
    # 실제 토큰 계산 전에 로그
    total_tokens = calculate_total_tokens(documents)
    
    # 250,000 토큰 제한에 맞춰 분할 처리 여부 결정
    if total_tokens > max_tokens:
        if logger:
            logger.update_detail(f"토큰 제한({max_tokens:,}) 초과. 배치 처리 시작 (총 {total_tokens:,} 토큰)", "info")
        
        # 안전한 배치 크기 (실제 토큰 기준) - 토큰 활용률 극대화
        safe_max_tokens = 240000  # 24만 토큰으로 증가 (OpenAI 제한의 80%)
        
        vectorstore = None
        current_batch = []
        current_tokens = 0
        batch_num = 1
                
        for i, doc in enumerate(documents):
            doc_tokens = count_actual_tokens(doc.page_content)
            
            # 단일 문서가 배치 크기를 크게 초과하는 경우에만 경고 및 건너뛰기 (50% 여유를 두어 중요한 문서 손실 방지)
            if doc_tokens > safe_max_tokens * 1.5:  # 현재는 30만 토큰(200,000 * 1.5)
                if logger:
                    logger.update_detail(f"문서 {i+1}이 너무 큽니다 ({doc_tokens:,} 토큰). 건너뜀", "warning")
                continue
            
            # 진행상황 표시 (25개마다) - 빈도 줄임
            if (i + 1) % 25 == 0:
                if logger:
                    logger.update_detail(f"벡터 임베딩 생성: {i+1}/{len(documents)} (배치 {batch_num})", "info")
            
            # 현재 배치에 추가해도 안전한지 확인
            if current_tokens + doc_tokens <= safe_max_tokens:
                current_batch.append(doc)
                current_tokens += doc_tokens
            else:
                # 현재 배치 처리
                if current_batch:
                    try:
                        if logger:
                            logger.update_detail(f"벡터스토어 배치 {batch_num} 생성 중... ({len(current_batch)}개 문서)", "info")
                        
                        if vectorstore is None:
                            vectorstore = FAISS.from_documents(current_batch, embedding_model)
                        else:
                            batch_vectorstore = FAISS.from_documents(current_batch, embedding_model)
                            vectorstore.merge_from(batch_vectorstore)
                        
                        if logger:
                            logger.update_detail(f"벡터스토어 배치 {batch_num} 완료", "success")
                        batch_num += 1
                        
                    except Exception as e:
                        if logger:
                            logger.update_detail(f"배치 {batch_num} 처리 실패: {e}", "error")
                        
                        # 더 작은 배치로 재시도
                        if len(current_batch) > 1:
                            if logger:
                                logger.update_detail(f"배치를 절반으로 나누어 재시도", "info")
                            
                            # 배치를 절반으로 나누어 재시도
                            mid = len(current_batch) // 2
                            for sub_batch in [current_batch[:mid], current_batch[mid:]]:
                                if sub_batch:
                                    try:
                                        if vectorstore is None:
                                            vectorstore = FAISS.from_documents(sub_batch, embedding_model)
                                        else:
                                            batch_vectorstore = FAISS.from_documents(sub_batch, embedding_model)
                                            vectorstore.merge_from(batch_vectorstore)
                                    except Exception as e2:
                                        if logger:
                                            logger.update_detail(f"소배치 실패: {e2}", "error")
                                        continue
                        else:
                            # 단일 문서도 실패하는 경우 건너뛰기
                            if logger:
                                logger.update_detail(f"단일 문서 처리 실패, 건너뜀: {e}", "warning")
                
                # 새 배치 시작
                current_batch = [doc]
                current_tokens = doc_tokens
        
        # 마지막 배치 처리
        if current_batch:
            try:
                if logger:
                    logger.update_detail(f"마지막 벡터스토어 배치 {batch_num} 생성 중... ({len(current_batch)}개 문서)", "info")
                
                if vectorstore is None:
                    vectorstore = FAISS.from_documents(current_batch, embedding_model)
                else:
                    batch_vectorstore = FAISS.from_documents(current_batch, embedding_model)
                    vectorstore.merge_from(batch_vectorstore)
                
                if logger:
                    logger.update_detail(f"마지막 벡터스토어 배치 {batch_num} 완료", "success")
                
            except Exception as e:
                if logger:
                    logger.update_detail(f"마지막 배치 처리 실패: {e}", "error")
                
                # 마지막 배치도 분할 시도
                if len(current_batch) > 1:
                    mid = len(current_batch) // 2
                    for sub_batch in [current_batch[:mid], current_batch[mid:]]:
                        if sub_batch:
                            try:
                                if vectorstore is None:
                                    vectorstore = FAISS.from_documents(sub_batch, embedding_model)
                                else:
                                    batch_vectorstore = FAISS.from_documents(sub_batch, embedding_model)
                                    vectorstore.merge_from(batch_vectorstore)
                                if logger:
                                    logger.update_detail(f"소배치 완료: {len(sub_batch)}개 문서", "success")
                            except:
                                continue
        
        if vectorstore is None:
            raise Exception("모든 문서 처리에 실패했습니다.")
        
        return vectorstore
    
    else:
        # 토큰 제한 내에 있으면 직접 처리
        try:
            if logger:
                logger.update_detail(f"토큰 제한 내 직접 처리 ({total_tokens:,} 토큰)", "info")
            return FAISS.from_documents(documents, embedding_model)
        except Exception as e:
            if logger:
                logger.update_detail(f"직접 처리 실패: {e}", "error")
            if "max_tokens_per_request" in str(e):
                # 강제 분할 처리로 재귀 호출 - 적당한 토큰 제한 적용
                return create_vectorstore_with_token_limit(documents, embedding_model, max_tokens=150000, logger=logger)
            else:
                raise e

# OpenAI 임베딩 모델 초기화
def create_embedding_model():
    """OpenAI 임베딩 모델을 생성합니다.
    
    토큰 제한 오류 시 배치 크기를 단계적으로 줄여가며 재시도합니다.
    - 1차: 1000개 배치 (빠른 처리)
    - 2차: 500개 배치 (안전)
    """
    
    try:
        embedding_model = OpenAIEmbeddings(
            model=OPENAI_EMBEDDING_MODEL,
            chunk_size=EMBEDDING_BATCH_SIZE_INITIAL,
            max_retries=OPENAI_MAX_RETRIES
        )
        return embedding_model
    except Exception as e:
        if "max_tokens_per_request" in str(e):
            # 로그 출력 제거 - 이미 상위 함수에서 처리됨
            return OpenAIEmbeddings(
                model=OPENAI_EMBEDDING_MODEL,
                chunk_size=EMBEDDING_BATCH_SIZE_RETRY,
                max_retries=OPENAI_MAX_RETRIES
            )
        else:
            raise e

# 텍스트 분할기 초기화 - 하이브리드 방식 (사전 분할 + SemanticChunker)
def create_text_splitter():
    """하이브리드 텍스트 분할기를 생성합니다.
    
    1단계: RecursiveCharacterTextSplitter로 큰 청크들로 사전 분할
    2단계: 각 청크에 SemanticChunker 적용하여 의미적 분할
    """
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain_experimental.text_splitter import SemanticChunker
    
    # 1단계: 사전 분할기 (큰 청크로 나누기) - 법령 문서 최적화
    pre_splitter = RecursiveCharacterTextSplitter(
        chunk_size=35000,  # 35,000자로 증가 (더 큰 문맥에서 의미 분석)
        chunk_overlap=100, # 100자로 최소화 (조문 연결부만 보존)
        length_function=len,
        separators=[
            "\n\n제", "\n\n조", "\n\n항", "\n\n호",  # 법령 구조 우선 (제/조/항/호)
            "\n\n\n\n", "\n\n\n", "\n\n", "\n",      # 일반 구조
            ".", ")", " ", ""                          # 최후 수단
        ]
    )
    
    return pre_splitter  # 우선 사전 분할기만 반환

def apply_semantic_chunking(text_chunk: str, embedding_model) -> list:
    """개별 텍스트 청크에 SemanticChunker를 적용합니다."""
    try:
        # 텍스트가 너무 짧으면 SemanticChunker 건너뛰기 (더 큰 사전 청크에 맞게 조정)
        if len(text_chunk) < 2000:
            return [text_chunk]
                
        # SemanticChunker 생성 및 적용
        semantic_splitter = SemanticChunker(
            embeddings=embedding_model,
            breakpoint_threshold_type="percentile",  # 다른 옵션: "standard_deviation", "interquartile"
            breakpoint_threshold_amount=90,  # 더욱 큰 의미 단위로 분할 (90%)
            buffer_size=2,  # 더 큰 청크에서는 버퍼 크기 증가로 정교한 분석
            sentence_split_regex=r'(?<=[.!?])\s+',  # 문장 분할 정규식 명시
            add_start_index=True  # 원본 텍스트에서의 시작 위치 추가
        )
        
        # 의미적 분할 수행
        semantic_chunks = semantic_splitter.split_text(text_chunk)
        
        # 빈 청크 제거 및 최소 길이 필터링 (더욱 관대한 기준으로 조정)
        filtered_chunks = [chunk for chunk in semantic_chunks if len(chunk.strip()) >= 100]
        
        return filtered_chunks if filtered_chunks else [text_chunk]
        
    except Exception as e:
        # SemanticChunker 실패 시 고정 크기 분할로 폴백
        # 로그 출력 제거 - 정상적인 처리 과정의 일부
        fallback_splitter = RecursiveCharacterTextSplitter(
            chunk_size=8000,   # 적절한 중간 크기 (25,000의 1/3 정도)
            chunk_overlap=100, # 1단계와 동일한 겹침
            length_function=len,
            separators=[
                "\n\n제", "\n\n조", "\n\n항", "\n\n호",  # 법령 구조 유지
                "\n\n\n", "\n\n", "\n", ".", " ", ""
            ]
        )
        return fallback_splitter.split_text(text_chunk)

def remove_duplicate_chunks(chunks: list) -> list:
    """중복되거나 매우 유사한 청크를 제거합니다. (법령 문서 최적화)"""
    if not chunks:
        return chunks
        
    # 법령 문서에서는 더 엄격한 중복 기준 적용 (65% 이상 유사하면 중복)
    similarity_threshold = 0.65
    unique_chunks = []
    
    for i, chunk in enumerate(chunks):
        is_duplicate = False
        chunk_text = chunk.strip()
        
        # 너무 짧은 청크는 제외
        if len(chunk_text) < 100:
            continue
            
        chunk_words = set(chunk_text.split())
        
        # 이미 추가된 청크들과 비교
        for j, existing_chunk in enumerate(unique_chunks):
            existing_text = existing_chunk.strip()
            existing_words = set(existing_text.split())
            
            # 두 청크 모두 충분한 길이인 경우에만 비교
            if len(chunk_words) > 10 and len(existing_words) > 10:
                # Jaccard 유사도 계산 (더 정확)
                intersection = len(chunk_words.intersection(existing_words))
                union = len(chunk_words.union(existing_words))
                jaccard_similarity = intersection / union if union > 0 else 0
                
                if jaccard_similarity >= similarity_threshold:
                    is_duplicate = True
                    break
        
        if not is_duplicate:
            unique_chunks.append(chunk)
    
    removed_count = len(chunks) - len(unique_chunks)
    
    return unique_chunks

# 문서 읽기 함수들
def read_docx(path):
    """DOCX 문서를 읽어서 텍스트를 추출합니다."""
    doc = Document(path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    text = latex_to_text(text)
    return text

def read_pdf(path):
    """PDF 문서를 읽어서 전체 텍스트를 추출합니다."""
    full_text = ""
    with pdfplumber.open(path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            try:
                page_text = page.extract_text()
                if page_text:
                    page_text = latex_to_text(page_text)
                    page_text = normalize_text(page_text)
                    full_text += page_text + "\n"
            except Exception as e:
                # 로그 출력 제거 - 정상적인 에러 처리
                continue
    return full_text

def read_txt(path):
    """TXT 파일을 읽어서 텍스트를 추출합니다."""
    with open(path, 'r', encoding='utf-8') as f:
        text = f.read()
    text = latex_to_text(text)
    return text

# 텍스트 정규화
def normalize_text(text):
    """텍스트 정규화 (한글-영문 공백, 수식 기호 보존)"""
    # 한글-영문 공백 정리
    text = re.sub(r'([가-힣])([A-Za-z0-9])', r'\1 \2', text)
    text = re.sub(r'([A-Za-z0-9])([가-힣])', r'\1 \2', text)

    # 수식 기호 보존: √ ± ≈ ∞ × ÷ π ² ³ ^ / = % 등
    math_symbols = "√±≈∞×÷π²³^=/%"

    # 더 관대한 정규화: 한글, 영문, 숫자, 공백, 문장부호, 수식 기호만 유지
    # 한글: 가-힣
    # 영문: A-Za-z
    # 숫자: 0-9
    # 공백: \s
    # 문장부호: .,!?;:-()[]{}
    # 수식 기호: √±≈∞×÷π²³^=/%
    # 추가 허용 문자: + (플러스 기호)
    
    allowed_pattern = r'[가-힣A-Za-z0-9\s.,!?;:\-\(\)\[\]\{\}' + re.escape(math_symbols + '+') + r']'
    text = re.sub(rf'[^{allowed_pattern}]', ' ', text)

    # 연속 공백 정리
    text = re.sub(r'\s+', ' ', text)

    return text.strip()

# 파일명 정규화 함수
def normalize_filename(file_name):
    """파일명을 정규화하여 분류에 사용합니다."""
    # 파일 확장자 제거
    name_without_ext = os.path.splitext(file_name)[0]
    
    # 특수문자 제거 (하이픈, 언더스코어, 플러스 등은 공백으로 변환)
    # 대괄호는 제거하되, 키워드가 포함된 부분은 보존
    normalized = re.sub(r'[\[\]\(\)\+\-\_]', ' ', name_without_ext)
    
    # 한글이 분해되지 않도록 NFC로 정규화
    normalized = unicodedata.normalize('NFC', normalized)
    
    # 키워드가 포함된 부분이 공백으로 분리되었는지 확인하고 복구
    if ' 전력 ' in normalized or normalized.startswith('전력 ') or normalized.endswith(' 전력'):
        # 전력 키워드 주변의 공백을 제거하여 복구
        normalized = re.sub(r'\s+전력\s+', '전력', normalized)
        normalized = re.sub(r'^전력\s+', '전력', normalized)
        normalized = re.sub(r'\s+전력$', '전력', normalized)
    
    if ' 도시가스 ' in normalized or normalized.startswith('도시가스 ') or normalized.endswith(' 도시가스'):
        # 도시가스 키워드 주변의 공백을 제거하여 복구
        normalized = re.sub(r'\s+도시가스\s+', '도시가스', normalized)
        normalized = re.sub(r'^도시가스\s+', '도시가스', normalized)
        normalized = re.sub(r'\s+도시가스$', '도시가스', normalized)
    
    # 연속 공백을 단일 공백으로 변환
    normalized = re.sub(r'\s+', ' ', normalized)
    
    # 앞뒤 공백 제거
    normalized = normalized.strip()
    
    return normalized

# 파일 분류 함수
def classify_file(file_name):
    """파일명에 따라 분류를 결정합니다."""
    # 파일명 정규화
    normalized_name = normalize_filename(file_name)
    
    # 키워드 검색 결과 (여러 방법으로 검색)
    gas_found = '도시가스' in normalized_name
    power_found = '전력' in normalized_name
    
    # 대안 검색 방법 (인코딩 문제 해결을 위해)
    gas_found_alt = normalized_name.find('도시가스') != -1
    power_found_alt = normalized_name.find('전력') != -1
    
    # 최종 결과 결정 (둘 중 하나라도 True면 True)
    gas_found = gas_found or gas_found_alt
    power_found = power_found or power_found_alt
    
    # 분류 결정
    if gas_found:
        return 'gas'
    elif power_found:
        return 'power'
    else:
        # 원본 파일명으로 재시도
        original_gas_found = '도시가스' in file_name
        original_power_found = '전력' in file_name
        
        if original_gas_found:
            return 'gas'
        elif original_power_found:
            return 'power'
        else:
            # 마지막 안전장치: 파일명의 모든 부분을 개별적으로 확인
            file_parts = re.split(r'[_\-\s\[\]\(\)]', file_name)
            
            for part in file_parts:
                if '도시가스' in part:
                    return 'gas'
                elif '전력' in part:
                    return 'power'
            
            return 'other'

def get_index_dir(category):
    """분류에 따른 인덱스 디렉토리를 반환합니다."""
    if category == 'gas':
        return GAS_INDEX_DIR
    elif category == 'power':
        return POWER_INDEX_DIR
    else:
        return OTHER_INDEX_DIR

def extract_metadata_from_filename(filename):
    base_name = os.path.splitext(filename)[0]
    parts = re.split(r'[_\-\s\+\(\)]', base_name)  # +와 ()도 구분자로 추가

    # 버전 추출 개선: 2024, 24.01.01, 2024.7.1 등 다양한 형식 지원
    version = None
    for part in parts:
        # 4자리 연도 포함 날짜 (2024.7.1) - 우선순위 1
        if re.match(r'20\d{2}\.\d{1,2}\.\d{1,2}', part):
            version = part  # 전체 날짜 보존
            break
        # 2자리 연도 포함 날짜 (24.01.01) - 우선순위 2
        elif re.match(r'\d{2}\.\d{1,2}\.\d{1,2}', part):
            year = part.split('.')[0]
            if int(year) >= 20:  # 20년 이후
                version = f"20{part}"  # 2024.01.01 형태로 변환
            break
        # 4자리 연도만 (2024) - 우선순위 3
        elif re.match(r'20\d{2}$', part):
            version = part
            break
    
    region_list = ['서울', '부산', '대구', '광주', '인천', '대전', '울산','경기도', '강원도', '충청북도', '충청남도' ,'전라남도','전북특별자치도', '경상남도', '경상북도']
    region = next((p for p in parts if p in region_list), None)
    organization_map = {'도시가스': '도시가스', '전력': '전력'}
    organization = next((p for p in organization_map if p in parts), "기타")

    return {
        "version": version,
        "region": region,
        "organization": organization,
        "title": base_name
    }

# 문서 처리 및 청킹
def process_document(file_path, logger=None):
    """문서를 읽고 SemanticChunker를 사용하여 의미적 청킹을 수행합니다.
    
    SemanticChunker는 문서의 의미를 고려하여 자동으로 적절한 크기의 청크로 분할합니다.
    """
    filename = os.path.basename(file_path)
    
    if logger:
        logger.update_detail(f"📖 텍스트 추출 중: {filename}")
    
    # 문서 읽기
    if file_path.endswith(".pdf"):
        text = read_pdf(file_path)
    elif file_path.endswith(".docx"):
        text = read_docx(file_path)
    elif file_path.endswith(".txt"):
        text = read_txt(file_path)
    else:
        # 로그 출력 제거 - 정상적인 에러 처리
        return []
    
    
    # 텍스트가 비어있으면 처리 중단
    if not text:
        if logger:
            logger.update_detail(f"❌ {filename}: 텍스트 추출 실패", "error")
        return []
    
    if logger:
        logger.update_detail(f"✅ 텍스트 추출 완료: {len(text):,}자")
    
    # 텍스트 정규화
    original_text = text
    text = normalize_text(text)
    
    # 정규화 과정에서 텍스트가 너무 많이 제거되었는지 확인
    if len(text) < len(original_text) * 0.1:  # 90% 이상 제거된 경우
        # 정규화를 건너뛰고 원본 텍스트 사용
        text = original_text
        
    # 문서 크기 기준 설정
    LARGE_DOCUMENT_THRESHOLD = 60000  # 6만자 이상은 대용량 문서로 판단
    
    if len(text) <= LARGE_DOCUMENT_THRESHOLD:
        # 작은 문서: SemanticChunker 직접 사용
        if logger:
            logger.update_detail(f"🧠 의미적 청킹 중...")
        
        embedding_model = create_embedding_model()
        
        try:
            chunks = apply_semantic_chunking(text, embedding_model)
        except Exception as e:
            # 실패 시 간단한 고정 크기 분할
            simple_splitter = RecursiveCharacterTextSplitter(
                chunk_size=3000,
                chunk_overlap=100,
                length_function=len,
                separators=[
                    "\n\n제", "\n\n조", "\n\n항", "\n\n호",  # 법령 구조 유지
                    "\n\n\n", "\n\n", "\n", ".", " ", ""
                ]
            )
            chunks = simple_splitter.split_text(text)
            
    else:
        # 대용량 문서: 하이브리드 방식 (사전 분할 + SemanticChunker)
        if logger:
            logger.update_detail(f"🧠 대용량 문서 청킹 중...")
        
        text_splitter = create_text_splitter()
        pre_chunks = text_splitter.split_text(text)
        
        # 2단계: 각 사전 청크에 SemanticChunker 적용
        all_chunks = []
        embedding_model = create_embedding_model()  # SemanticChunker용 임베딩 모델
        
        for i, pre_chunk in enumerate(pre_chunks):
            if logger and (i + 1) % 5 == 0:  # 5개마다 진행상황 표시
                logger.update_detail(f"🔄 의미적 청킹 진행: {i+1}/{len(pre_chunks)}")
            
            try:
                semantic_chunks = apply_semantic_chunking(pre_chunk, embedding_model)
                all_chunks.extend(semantic_chunks)
            except Exception as e:
                # 실패 시 원본 사전 청크를 더 작게 나누어 추가
                fallback_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=8000,   # apply_semantic_chunking과 동일한 설정
                    chunk_overlap=100, # 1단계와 동일한 겹침
                    length_function=len,
                    separators=[
                        "\n\n제", "\n\n조", "\n\n항", "\n\n호",  # 법령 구조 유지
                        "\n\n\n", "\n\n", "\n", ".", " ", ""
                    ]
                )
                fallback_chunks = fallback_splitter.split_text(pre_chunk)
                all_chunks.extend(fallback_chunks)
        
        chunks = all_chunks
    
    # 3단계: 중복 청크 제거 (법령 문서 정확성 확보)
    chunks = remove_duplicate_chunks(chunks)
    
    if logger:
        logger.update_detail(f"✅ 청킹 완료: {len(chunks)}개 문서 생성")
    
    # 파일명에서 추가 메타데이터 추출
    additional_metadata = extract_metadata_from_filename(filename)
    
    # 메타데이터가 포함된 LangChain 문서로 변환
    documents = []
    for i, chunk in enumerate(chunks):
        if len(chunk.strip()) > 100:  # 최소 길이 필터 (100자 이상)
            # 기본 메타데이터와 추가 메타데이터를 결합
            metadata = {
                "source": filename,
                "chunk_id": i,
                "file_path": file_path,
                "version": additional_metadata["version"],
                "region": additional_metadata["region"],
                "organization": additional_metadata["organization"],
                "title": additional_metadata["title"]
            }
            doc = LCDocument(
                page_content=chunk,
                metadata=metadata
            )
            documents.append(doc)
    
    return documents

# 벡터 인덱스 구축
def build_vector_index_from_uploaded_files(uploaded_files):
    """업로드된 파일들로부터 벡터 인덱스를 구축합니다."""
   
    if not uploaded_files:
        st.warning("업로드된 파일이 없습니다.")
        return False
    
    # 통합 로거 생성
    logger = UnifiedLogger()
    logger.start("AI 문서 학습 시작")
    
    try:
        # 문서 저장 디렉토리 생성
        docs_dir = Path(__file__).parent.parent / "vectordb" / "docs"
        docs_dir.mkdir(parents=True, exist_ok=True)
        
        # 임베딩 모델 초기화
        logger.update_main("임베딩 모델 초기화", "진행중", "🔧")
        embedding_model = create_embedding_model()
        
        # 분류별로 문서 그룹화
        gas_documents = []
        power_documents = []
        other_documents = []
        
        total_files = len(uploaded_files)
        logger.update_main(f"파일 처리 ({total_files}개)", "진행중", "📝")
        
        # 파일별로 처리 및 분류
        for i, uploaded_file in enumerate(uploaded_files, 1):
            logger.update_detail(f"[{i}/{total_files}] 처리 중: {uploaded_file.name}")
            
            # 실제 파일을 docs 디렉토리에 저장
            file_path = docs_dir / uploaded_file.name
            
            # 이미 동일한 파일명이 존재하는지 확인
            if file_path.exists():
                logger.error(f"'{uploaded_file.name}' 파일이 이미 존재합니다.")
                return False
            
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getvalue())
            
            try:
                # 파일 분류
                category = classify_file(uploaded_file.name)
                
                documents = process_document(str(file_path), logger)
                
                # 분류별로 문서 추가
                if category == 'gas':
                    gas_documents.extend(documents)
                elif category == 'power':
                    power_documents.extend(documents)
                else:
                    other_documents.extend(documents)
                
                logger.update_detail(f"✅ {uploaded_file.name} 처리 완료 ({category} 분류, {len(documents)}개 문서)", "success")
                
                # 메모리 최적화
                del documents
                gc.collect()
                
            except Exception as e:
                logger.error(f"파일 처리 오류: {e}")
                # 오류 발생 시 저장된 파일 삭제
                if file_path.exists():
                    file_path.unlink()
                return False
        
        # 분류별로 인덱스 생성
        logger.update_main("벡터 인덱스 생성", "진행중", "🔧")
        
        categories = [
            ('gas', gas_documents, 'Gas'),
            ('power', power_documents, 'Power'),
            ('other', other_documents, 'Other')
        ]
        
        success_count = 0
        for category, documents, category_name in categories:
            if len(documents) == 0:
                logger.update_detail(f"{category_name} 분류: 문서 없음, 건너뜀")
                continue
                
            logger.update_detail(f"{category_name} 인덱스 생성 중... (문서 수: {len(documents)}개)")
            
            # 배치 단위로 임베딩 처리
            try:
                vectorstore = create_vectorstore_with_token_limit(documents, embedding_model, logger=logger)
            except Exception as e:
                if "max_tokens_per_request" in str(e):
                    try:
                        logger.update_detail(f"{category_name}: 토큰 제한으로 배치 크기 조정 (재시도 1/2)", "warning")
                        medium_embedding_model = OpenAIEmbeddings(
                            model=OPENAI_EMBEDDING_MODEL,
                            chunk_size=EMBEDDING_BATCH_SIZE_RETRY,
                            max_retries=OPENAI_MAX_RETRIES
                        )
                        vectorstore = create_vectorstore_with_token_limit(documents, medium_embedding_model, logger=logger)
                    except Exception as e2:
                        if "max_tokens_per_request" in str(e2):
                            logger.update_detail(f"{category_name}: 토큰 제한으로 배치 크기 재조정 (재시도 2/2)", "warning")
                            small_embedding_model = OpenAIEmbeddings(
                                model=OPENAI_EMBEDDING_MODEL,
                                chunk_size=EMBEDDING_BATCH_SIZE_FINAL,
                                max_retries=OPENAI_MAX_RETRIES
                            )
                            vectorstore = create_vectorstore_with_token_limit(documents, small_embedding_model, logger=logger)
                        else:
                            raise e2
                else:
                    raise e

            # 인덱스 저장
            index_dir = get_index_dir(category)
            try:
                vectorstore.save_local(str(index_dir))
                logger.update_detail(f"{category_name} 인덱스 저장 완료", "success")
                success_count += 1
            except Exception as e:
                logger.error(f"{category_name} 인덱스 저장 오류: {e}")
                return False
        
        if success_count > 0:
            logger.success(f"🎉 {success_count}개 분류별 인덱스 생성 완료")
            return True
        else:
            logger.error("❌ 인덱스 생성에 실패했습니다.")
            return False
            
    except Exception as e:
        logger.error(f"예상치 못한 오류: {e}")
        return False

def reload_backend_indexes():
    """백엔드의 인덱스를 다시 로드합니다."""
    # 통합 로거 생성
    logger = UnifiedLogger()
    logger.start("백엔드 인덱스 리로드")
    
    try:
        logger.update_detail("API 서버에 리로드 요청 중...")
        response = requests.post(RELOAD_API_URL)
        response.raise_for_status()
        
        result = response.json()
        if result.get("success"):
            logger.success("✅ 백엔드 인덱스 리로드 완료!")
            return True
        else:
            logger.error(f"백엔드 인덱스 리로드 실패: {result.get('message', '알 수 없는 오류')}")
            return False
    except Exception as e:
        logger.error(f"백엔드 인덱스 리로드 중 오류: {str(e)}")
        st.warning("⚠️ FastAPI 서버가 실행 중인지 확인해주세요.")
        return False

def latex_to_text(text):
    """
    LaTeX 수식을 사람이 읽는 텍스트 수식으로 변환
    예: \frac{a}{b} → (a) / (b)
    """
    # \frac 변환 함수
    def frac_repl(match):
        return f"({match.group(1)}) / ({match.group(2)})"

    # LaTeX 블록(\[...\], $$...$$, $...$)을 찾아서 변환
    def latex_block_repl(match):
        latex_expr = match.group(1)
        # \left와 \right 제거 (괄호 크기 조정 명령어) - 더 포괄적으로 처리
        latex_expr = re.sub(r'\\left\s*\(', '(', latex_expr)
        latex_expr = re.sub(r'\\right\s*\)', ')', latex_expr)
        latex_expr = re.sub(r'\\left\s*\[', '[', latex_expr)
        latex_expr = re.sub(r'\\right\s*\]', ']', latex_expr)
        latex_expr = re.sub(r'\\left\s*\\{', '{', latex_expr)
        latex_expr = re.sub(r'\\right\s*\\}', '}', latex_expr)
        # \frac 변환
        latex_expr = re.sub(r'\\frac\{(.+?)\}\{(.+?)\}', frac_repl, latex_expr)
        # \times 변환
        latex_expr = latex_expr.replace(r'\times', '×')
        # 중괄호 제거
        latex_expr = latex_expr.replace('{', '').replace('}', '')
        return latex_expr

    # \[ ... \] 블록 변환
    text = re.sub(r'\\\[(.*?)\\\]', lambda m: latex_block_repl(m), text, flags=re.DOTALL)
    # $$ ... $$ 블록 변환
    text = re.sub(r'\$\$(.*?)\$\$', lambda m: latex_block_repl(m), text, flags=re.DOTALL)
    # $ ... $ 블록 변환
    text = re.sub(r'\$(.*?)\$', lambda m: latex_block_repl(m), text, flags=re.DOTALL)

    # 인라인 변환 (혹시 남아있을 경우)
    # \left와 \right 제거 - 더 포괄적으로 처리
    text = re.sub(r'\\left\s*\(', '(', text)
    text = re.sub(r'\\right\s*\)', ')', text)
    text = re.sub(r'\\left\s*\[', '[', text)
    text = re.sub(r'\\right\s*\]', ']', text)
    text = re.sub(r'\\left\s*\\{', '{', text)
    text = re.sub(r'\\right\s*\\}', '}', text)
    # \frac 변환
    text = re.sub(r'\\frac\{(.+?)\}\{(.+?)\}', frac_repl, text)
    text = text.replace(r'\times', '×')
    text = text.replace('{', '').replace('}', '')

    return text

def initialize_session_state():
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "streaming" not in st.session_state:
        st.session_state.streaming = False
    if "current_page" not in st.session_state:
        st.session_state.current_page = "chat"
    if "api_processing" not in st.session_state:
        st.session_state.api_processing = False




def highlight_relevant_content(content: str, search_terms: list) -> str:
    """문서 내용에서 관련 부분을 하이라이트합니다."""
    if not search_terms:
        return content
    
    highlighted_content = content
    
    # 각 검색어에 대해 하이라이트 적용
    for term in search_terms:
        if term and len(term.strip()) > 1:
            # HTML 하이라이트 적용
            highlighted_content = re.sub(
                f'({re.escape(term)})',
                r'<mark style="background-color: #ffeb3b; padding: 1px 2px;">\1</mark>',
                highlighted_content,
                flags=re.IGNORECASE
            )
    
    return highlighted_content

def extract_keywords_from_recent_chat() -> list:
    """최근 채팅에서 키워드를 추출합니다."""
    if not st.session_state.messages:
        return []
    
    # 최근 사용자 질문에서 키워드 추출
    recent_user_messages = [
        msg["content"] for msg in st.session_state.messages[-5:] 
        if msg["role"] == "user"
    ]
    
    keywords = []
    for message in recent_user_messages:
        # 간단한 키워드 추출 (한글 2글자 이상)
        words = re.findall(r'[가-힣]{2,}', message)
        keywords.extend(words)
    
    # 중복 제거 및 길이 정렬 (긴 키워드 우선)
    unique_keywords = list(set(keywords))
    unique_keywords.sort(key=len, reverse=True)
    
    return unique_keywords[:10]  # 상위 10개만 사용



def extract_document_links(content: str) -> list:
    """메시지 내용에서 문서 링크를 추출합니다."""
    pattern = r'\[DOCUMENT:([^\]]+)\]'
    matches = re.findall(pattern, content)
    return matches

def find_actual_file_path(docs_dir: Path, filename: str) -> Path:
    """파일명을 정확히 찾기 위해 다양한 패턴으로 검색합니다."""
    print(f"DEBUG: 찾는 파일명: {filename}")
    print(f"DEBUG: 검색 디렉토리: {docs_dir}")
    
    # 디렉토리의 모든 파일 나열
    try:
        all_files = list(docs_dir.iterdir())
        print(f"DEBUG: 디렉토리의 모든 파일: {[f.name for f in all_files if f.is_file()]}")
    except Exception as e:
        print(f"DEBUG: 디렉토리 읽기 오류: {e}")
    
    # 1. 정확한 파일명으로 찾기
    file_path = docs_dir / filename
    print(f"DEBUG: 정확한 파일명으로 검색: {file_path}")
    if file_path.exists():
        print(f"DEBUG: 파일 발견! {file_path}")
        return file_path
    
    # 1-1. 파일명에 대괄호가 있는 경우, 그대로 검색
    if '[' in filename and ']' in filename:
        # 이미 대괄호가 있는 파일명이므로 그대로 사용
        file_path = docs_dir / filename
        print(f"DEBUG: 대괄호 포함 파일명으로 검색: {file_path}")
        if file_path.exists():
            print(f"DEBUG: 대괄호 포함 파일명으로 파일 발견! {file_path}")
            return file_path
    
    # 2. 대괄호가 빠진 경우를 고려해서 찾기
    if not filename.startswith('['):
        print(f"DEBUG: 대괄호가 없는 파일명, 키워드 검색 시작")
        keywords = ['도시가스', '전력', '기타']
        for keyword in keywords:
            if keyword in filename:
                bracketed_filename = f"[{keyword}]{filename}"
                file_path = docs_dir / bracketed_filename
                print(f"DEBUG: 키워드 {keyword}로 검색: {file_path}")
                if file_path.exists():
                    print(f"DEBUG: 키워드 검색으로 파일 발견! {file_path}")
                    return file_path
    
    # 3. 대괄호가 있는 경우, 대괄호를 제거한 형태로도 찾기
    if filename.startswith('[') and ']' in filename:
        clean_filename = filename.split(']', 1)[1] if ']' in filename else filename
        file_path = docs_dir / clean_filename
        if file_path.exists():
            return file_path
    
    # 4. 지역명 우선 매칭 검색
    try:
        print(f"DEBUG: 지역명 우선 매칭 검색 시작")
        
        # 요청 파일명에서 지역명 추출
        def extract_region_from_filename(fname):
            """파일명에서 지역명을 추출합니다."""
            regions = ['서울특별시', '강원도', '경기도', '경상북도', '전라남도', '충청북도', '부산', '서울', '부산광역시', '대구광역시', '인천광역시', '광주광역시', '대전광역시', '울산광역시', '세종특별자치시', '경기도', '강원도', '충청북도', '충청남도', '전라북도', '전라남도', '경상북도', '경상남도', '제주특별자치도']
            for region in regions:
                if region in fname:
                    return region
            return None
        
        requested_region = extract_region_from_filename(filename)
        print(f"DEBUG: 요청 파일명에서 추출한 지역: {requested_region}")
        
        # 지역명이 있는 경우, 정확한 지역 파일만 찾기
        if requested_region:
            print(f"DEBUG: 지역명 '{requested_region}'으로 정확한 지역 파일 검색")
            for file in docs_dir.iterdir():
                if file.is_file():
                    file_region = extract_region_from_filename(file.name)
                    if file_region and requested_region in file_region:
                        # 추가로 파일 내용이 유사한지 확인
                        clean_request = filename.replace('[', '').replace(']', '').replace('+', ' ').lower()
                        clean_file = file.name.replace('[', '').replace(']', '').replace('+', ' ').lower()
                        
                        # 파일명의 주요 키워드들이 일치하는지 확인
                        request_keywords = set(clean_request.split())
                        file_keywords = set(clean_file.split())
                        
                        # 60% 이상 키워드가 일치하면 해당 파일로 선택
                        if len(request_keywords.intersection(file_keywords)) / len(request_keywords) >= 0.6:
                            print(f"DEBUG: 지역명 매칭으로 파일 발견! {file}")
                            return file
        
        # 지역명이 없거나 지역명 매칭에 실패한 경우, 부분 매칭으로 검색
        print(f"DEBUG: 부분 매칭 검색 시작")
        for file in docs_dir.iterdir():
            if file.is_file():
                print(f"DEBUG: 검사 중인 파일: {file.name}")
                
                # 파일명 정규화 (대괄호, 공백, 특수문자 처리)
                clean_request = filename.replace('[', '').replace(']', '').replace('+', ' ').lower().strip()
                clean_file = file.name.replace('[', '').replace(']', '').replace('+', ' ').lower().strip()
                
                print(f"DEBUG: 정규화된 요청: '{clean_request}'")
                print(f"DEBUG: 정규화된 파일: '{clean_file}'")
                
                # 1. 정확한 부분 문자열 매칭
                if clean_request in clean_file or clean_file in clean_request:
                    print(f"DEBUG: 부분 문자열 매칭으로 파일 발견! {file}")
                    return file
                
                # 2. 키워드 기반 매칭
                request_keywords = set(clean_request.split())
                file_keywords = set(clean_file.split())
                
                # 공통 키워드가 있는지 확인
                common_keywords = request_keywords.intersection(file_keywords)
                if common_keywords:
                    print(f"DEBUG: 공통 키워드 발견: {common_keywords}")
                    # 50% 이상 키워드가 일치하면 해당 파일로 선택
                    if len(common_keywords) / len(request_keywords) >= 0.5:
                        print(f"DEBUG: 키워드 매칭으로 파일 발견! {file}")
                        return file
                
                # 3. 특정 키워드가 포함된 경우 (전력, 도시가스 등)
                important_keywords = ['전력', '도시가스', '기타', '비용평가', '운영규정']
                for keyword in important_keywords:
                    if keyword in clean_request and keyword in clean_file:
                        print(f"DEBUG: 중요 키워드 '{keyword}' 매칭으로 파일 발견! {file}")
                        return file
                        
    except Exception as e:
        print(f"DEBUG: 검색 중 오류: {e}")
        pass
    
    return None



def display_chat_history():
    for idx, message in enumerate(st.session_state.messages):
        with st.chat_message(message["role"]):
            if message["role"] == "assistant":
                # [DOCUMENT:[filename]] 패턴을 다운로드 가능한 파일명으로 변환
                content = message["content"]
                
                def create_download_link(match):
                    filename = match.group(1)
                    # 파일 경로 생성 (업로드 경로와 일치하도록 수정)
                    docs_dir = Path(__file__).parent.parent / "vectordb" / "docs"
                    file_path = find_actual_file_path(docs_dir, filename)
                    
                    if file_path and file_path.exists():
                        # 파일을 읽어서 다운로드 가능한 링크 생성
                        try:
                            with open(file_path, "rb") as f:
                                file_data = f.read()
                            
                            # base64로 인코딩하여 다운로드 링크 생성
                            import base64
                            b64_data = base64.b64encode(file_data).decode()
                            file_ext = file_path.suffix.lower()
                            
                            # MIME 타입 설정
                            mime_types = {
                                '.pdf': 'application/pdf',
                                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                '.txt': 'text/plain'
                            }
                            mime_type = mime_types.get(file_ext, 'application/octet-stream')
                            
                            # 다운로드 링크 HTML 생성
                            href = f'<a href="data:{mime_type};base64,{b64_data}" download="{file_path.name}" style="color: #2d9bf0; text-decoration: underline; font-weight: bold;">📋 {filename}</a>'
                            return href
                        except Exception as e:
                            return f'📥 다운로드 (오류)'
                    else:
                        return f'📥 다운로드 (파일 없음)'
                
                pattern = r'\[DOCUMENT:(.*)\]'
                content = re.sub(pattern, create_download_link, content)
                
                # 메시지 내용 표시 (HTML 허용)
                st.markdown(content, unsafe_allow_html=True)
            else:
                st.markdown(message["content"])

def stream_response(response_text: str, loading_placeholder):
    message_placeholder = st.empty()
    full_response = ""
    loading_placeholder.empty()
    for char in response_text:
        full_response += char
        message_placeholder.markdown(full_response + "▌")
        time.sleep(0.01)
    message_placeholder.markdown(full_response)
    return full_response

def send_message(user_input: str):
    if not user_input:
        return
    
    # API 처리 상태 시작 (즉시 설정)
    st.session_state.api_processing = True
    
    # 애니메이션 상태 초기화
    if "animation_step" not in st.session_state:
        st.session_state.animation_step = 0
    
    # 사용자 메시지 표시
    with st.chat_message("user"):
        st.markdown(user_input)
    
    # 로딩 메시지를 위한 플레이스홀더 
    loading_placeholder = st.empty()
    
    # 메시지를 세션에 추가
    st.session_state.messages.append({"role": "user", "content": user_input})

    # API 요청 데이터 준비
    request_data = {
        "messages": [
            {"role": msg["role"], "content": msg["content"]}
            for msg in st.session_state.messages
        ]
    }
    
    try:
        # 연속된 애니메이션과 API 호출을 동시에 처리
        dot_patterns = ["", ".", "..", "...", "...."]
        
        # API 호출을 별도 스레드에서 실행
        import concurrent.futures
        import threading
        
        result_container = {"response": None, "error": None, "completed": False}
        
        def api_call():
            try:
                response = requests.post(API_URL, json=request_data)
                response.raise_for_status()
                result_container["response"] = response.json()["response"]
            except Exception as e:
                result_container["error"] = str(e)
            finally:
                result_container["completed"] = True
        
        # API 호출 스레드 시작
        api_thread = threading.Thread(target=api_call, daemon=True)
        api_thread.start()
        
        # 애니메이션 실행 (API 완료까지)
        cycle = 0
        while not result_container["completed"]:
            dots = dot_patterns[cycle % len(dot_patterns)]
            loading_placeholder.markdown(
                f"<div style='font-size: 1rem; color: #6b7280;'>🔍 AI가 응답을 생성하고 있습니다{dots}</div>",
                unsafe_allow_html=True
            )
            time.sleep(0.4)
            cycle += 1
        
        # API 호출 완료 대기
        api_thread.join()
        
        if result_container["error"]:
            raise Exception(result_container["error"])
        
        # 응답 처리
        assistant_response = result_container["response"]
        
        # 수식 변환 적용
        assistant_response = latex_to_text(assistant_response)
        
        # API 처리 완료 상태로 변경
        st.session_state.api_processing = False
        
        # 스트리밍 방식으로 응답 표시
        with st.chat_message("assistant"):
            streamed_response = stream_response(assistant_response, loading_placeholder)
            st.session_state.messages.append({"role": "assistant", "content": streamed_response})
        
    except Exception as e:
        # 오류 발생 시에도 API 처리 완료 상태로 변경
        st.session_state.api_processing = False
        loading_placeholder.empty()
        st.error(f"Error: {str(e)}")
    
    # 최종적으로 API 처리 상태 확실히 종료
    st.session_state.api_processing = False

def show_upload_page():
    """문서 업로드 페이지를 표시합니다."""
    # 내부 로고 영역 (채팅 화면과 동일한 스타일)
    st.markdown("""
        <div style="background-color:#ffffff; padding: 1rem 2rem; border-radius: 6px; margin-bottom: 1rem; display: flex; align-items: center; border: 1px solid #e5e7eb; box-shadow: 0 2px 4px rgba(0,0,0,0.04);">
            <img src="https://img.icons8.com/ios-filled/50/2d9bf0/document--v1.png" width="24px" style="margin-right: 10px;" />
            <span style="font-size: 1.3rem; font-weight: bold; color: #111827;">문서 업로드 </span>
        </div>
        <div style="color: #666666; font-size: 0.95rem; margin-bottom: 1.5rem;">
            문서를 업로드 해 AI를 학습시켜보세요.
        </div>
    """, unsafe_allow_html=True)

    st.markdown("""
        <div style="background-color: rgba(30, 41, 59, 0.9); padding: 1rem 1.5rem; border-radius: 10px; border: 1px solid #1f2937; color: #f3f4f6; font-size: 0.95rem; line-height: 1.7;">
        <strong style="font-size: 1.1rem; color: #ffffff;">📌 파일 분류 규칙</strong>
        <ul style="list-style-type: '📂 '; padding-left: 1.2em; margin: 0;">
            <li><b>도시가스</b> 키워드가 파일명에 포함 → <span style="color: #38bdf8;"><b>Gas</b></span> 분류</li>
            <li><b>전력</b> 키워드가 파일명에 포함 → <span style="color: #38bdf8;"><b>Power</b></span> 분류</li>
            <li>그 외 키워드일 경우 → <span style="color: #facc15;"><b>Other</b></span> 분류</li>
        </ul>
        </div>
    """, unsafe_allow_html=True)
    
    st.write("")  # 공백 추가
    
    uploaded_files = st.file_uploader(
        "문서 파일을 선택하세요",
        type=['pdf', 'docx', 'txt'],
        accept_multiple_files=True,
        help="여러 파일을 동시에 업로드할 수 있습니다.",
        label_visibility="collapsed"
    )
    # st.caption("PDF, DOCX, TXT 파일을 업로드 가능 합니다. (최대 200MB)")

    st.write("")  # 공백 추가

    if uploaded_files:
        file_data = []

        for file in uploaded_files:
            file_name = file.name
            file_size_kb = len(file.getvalue()) / 1024
            category = classify_file(file_name)

            warning_msg = ""
            if category == "other":
                warning_msg = "⚠️ 분류 불확실"

            file_data.append({
                "파일명": file_name,
                "크기": f"{file_size_kb:.1f} KB" if file_size_kb < 1024 else f"{file_size_kb/1024:.2f} MB",
                "분류": warning_msg if warning_msg else category
            })

        df = pd.DataFrame(file_data)
        st.markdown("""
            <div style="background-color:#ffffff; padding: 1rem 2rem; border-radius: 6px; margin-bottom: 1rem; display: flex; align-items: center; border: 1px solid #e5e7eb; box-shadow: 0 2px 4px rgba(0,0,0,0.04);">
                <img src="https://img.icons8.com/ios-filled/50/2d9bf0/database--v1.png" width="24px" style="margin-right: 10px;" />
                <span style="font-size: 1.3rem; font-weight: bold; color: #111827;">📄 업로드된 문서</span>
            </div>
        """, unsafe_allow_html=True)
        st.table(df)
 
    if st.button("▶️ AI 문서 학습 시작", type="primary", disabled=not uploaded_files):
        try:
            success = build_vector_index_from_uploaded_files(uploaded_files)
            if success:
                # 백엔드 인덱스 리로드
                reload_success = reload_backend_indexes()
                if reload_success:
                    st.success("🎉 모든 작업이 완료되었습니다! 이제 채팅 탭에서 질문할 수 있습니다.")
                else:
                    st.warning("⚠️ 인덱스는 생성되었지만 백엔드 리로드에 실패했습니다. FastAPI 서버를 재시작하거나 [학습 문서 정보 리로드] 버튼을 눌러 주세요.")
        except Exception as e:
            st.error(f"❌ 인덱스 생성 중 오류가 발생했습니다: {str(e)}")
            st.warning("API 서버가 실행되지 않았을 수 있습니다. FastAPI 서버를 먼저 실행해주세요.")
    
    st.write("")  # 공백 추가

    # 기존 학습 정보 헤더
    st.markdown("""
        <div style="background-color:#ffffff; padding: 1rem 2rem; border-radius: 6px; margin-bottom: 1rem; display: flex; align-items: center; border: 1px solid #e5e7eb; box-shadow: 0 2px 4px rgba(0,0,0,0.04);">
            <img src="https://img.icons8.com/ios-filled/50/2d9bf0/database--v1.png" width="24px" style="margin-right: 10px;" />
            <span style="font-size: 1.3rem; font-weight: bold; color: #111827;">학습 완료된 문서 정보</span>
        </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    
    # Gas 인덱스 카드
    with col1:
        if GAS_INDEX_DIR.exists() and any(GAS_INDEX_DIR.iterdir()):
            # 저장된 문서 수 계산
            docs_dir = Path(__file__).parent.parent / "vectordb" / "docs"
            gas_docs = [f for f in docs_dir.iterdir() if f.is_file() and classify_file(f.name) == 'gas'] if docs_dir.exists() else []
            doc_count = len(gas_docs)
            
            # 실제 마지막 업데이트 날짜 계산
            import datetime
            index_files = list(GAS_INDEX_DIR.iterdir())
            if index_files:
                latest_time = max(f.stat().st_mtime for f in index_files)
                last_update = datetime.datetime.fromtimestamp(latest_time).strftime("%Y-%m-%d")
            else:
                last_update = "알 수 없음"
            
            st.markdown(f"""
                <div style="background-color: #f8f9fa; padding: 1.5rem; border-radius: 10px; border: 1px solid #e9ecef; text-align: center;">
                    <div style="font-weight: bold; font-size: 1.1rem; color: #333; margin-bottom: 0.5rem;">✅ Gas 문서 학습 완료</div>
                    <div style="color: #666; font-size: 0.9rem; margin-bottom: 0.3rem;">학습 완료 문서 {doc_count}건</div>
                    <div style="color: #999; font-size: 0.8rem;">마지막 업데이트: {last_update}</div>
                </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown(f"""
                <div style="background-color: #f8f9fa; padding: 1.5rem; border-radius: 10px; border: 1px solid #e9ecef; text-align: center;">
                    <div style="font-weight: bold; font-size: 1.1rem; color: #333; margin-bottom: 0.5rem;">⚠️ Gas 문서 학습 미완료</div>
                    <div style="color: #666; font-size: 0.9rem; margin-bottom: 0.3rem;">학습 완료 문서 0건</div>
                    <div style="color: #999; font-size: 0.8rem;"></div>
                </div>
            """, unsafe_allow_html=True)
    
    # Power 인덱스 카드
    with col2:
        if POWER_INDEX_DIR.exists() and any(POWER_INDEX_DIR.iterdir()):
            # 저장된 문서 수 계산
            docs_dir = Path(__file__).parent.parent / "vectordb" / "docs"
            power_docs = [f for f in docs_dir.iterdir() if f.is_file() and classify_file(f.name) == 'power'] if docs_dir.exists() else []
            doc_count = len(power_docs)
            
            # 실제 마지막 업데이트 날짜 계산
            import datetime
            index_files = list(POWER_INDEX_DIR.iterdir())
            if index_files:
                latest_time = max(f.stat().st_mtime for f in index_files)
                last_update = datetime.datetime.fromtimestamp(latest_time).strftime("%Y-%m-%d")
            else:
                last_update = "알 수 없음"
            
            st.markdown(f"""
                <div style="background-color: #f8f9fa; padding: 1.5rem; border-radius: 10px; border: 1px solid #e9ecef; text-align: center;">
                    <div style="font-weight: bold; font-size: 1.1rem; color: #333; margin-bottom: 0.5rem;">✅ Power 문서 학습 완료</div>
                    <div style="color: #666; font-size: 0.9rem; margin-bottom: 0.3rem;">학습 완료 문서 {doc_count}건</div>
                    <div style="color: #999; font-size: 0.8rem;">마지막 업데이트: {last_update}</div>
                </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown(f"""
                <div style="background-color: #f8f9fa; padding: 1.5rem; border-radius: 10px; border: 1px solid #e9ecef; text-align: center;">
                    <div style="font-weight: bold; font-size: 1.1rem; color: #333; margin-bottom: 0.5rem;">⚠️ Power 문서 학습 미완료</div>
                    <div style="color: #666; font-size: 0.9rem; margin-bottom: 0.3rem;">학습 완료 문서 0건</div>
                    <div style="color: #999; font-size: 0.8rem;"></div>
                </div>
            """, unsafe_allow_html=True)
    
    # Other 인덱스 카드
    with col3:
        if OTHER_INDEX_DIR.exists() and any(OTHER_INDEX_DIR.iterdir()):
            # 저장된 문서 수 계산
            docs_dir = Path(__file__).parent.parent / "vectordb" / "docs"
            other_docs = [f for f in docs_dir.iterdir() if f.is_file() and classify_file(f.name) == 'other'] if docs_dir.exists() else []
            doc_count = len(other_docs)
            
            # 실제 마지막 업데이트 날짜 계산
            import datetime
            index_files = list(OTHER_INDEX_DIR.iterdir())
            if index_files:
                latest_time = max(f.stat().st_mtime for f in index_files)
                last_update = datetime.datetime.fromtimestamp(latest_time).strftime("%Y-%m-%d")
            else:
                last_update = "알 수 없음"
            
            st.markdown(f"""
                <div style="background-color: #f8f9fa; padding: 1.5rem; border-radius: 10px; border: 1px solid #e9ecef; text-align: center;">
                    <div style="font-weight: bold; font-size: 1.1rem; color: #333; margin-bottom: 0.5rem;">✅ Other 문서 학습 완료</div>
                    <div style="color: #666; font-size: 0.9rem; margin-bottom: 0.3rem;">학습 완료 문서 {doc_count}건</div>
                    <div style="color: #999; font-size: 0.8rem;">마지막 업데이트: {last_update}</div>
                </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown(f"""
                <div style="background-color: #f8f9fa; padding: 1.5rem; border-radius: 10px; border: 1px solid #e9ecef; text-align: center;">
                    <div style="font-weight: bold; font-size: 1.1rem; color: #333; margin-bottom: 0.5rem;">⚠️ Other 문서 학습 미완료</div>
                    <div style="color: #666; font-size: 0.9rem; margin-bottom: 0.3rem;">학습 완료 문서 0건</div>
                    <div style="color: #999; font-size: 0.8rem;"></div>
                </div>
            """, unsafe_allow_html=True)
    
    
    st.write("")  # 공백 추가
    
    # 저장된 문서 파일 목록 표시
    docs_dir = Path(__file__).parent.parent / "vectordb" / "docs"
    if docs_dir.exists() and any(docs_dir.iterdir()):
        # 숨김 파일 제외하고 파일 목록 가져오기
        files = [f for f in docs_dir.iterdir() if f.is_file() and not f.name.startswith('.')]
        if files:
            # st.info(f"📂 총 {len(files)}개 파일이 저장되어 있습니다:")
            
            # 파일 데이터를 표 형태로 준비
            file_data = []
            for file in sorted(files):
                file_size = file.stat().st_size
                file_size_kb = file_size / 1024
                category = classify_file(file.name)
                
                file_data.append({
                    "파일명": file.name,
                    "사이즈": f"{file_size_kb:.1f} KB" if file_size_kb < 1024 else f"{file_size_kb/1024:.2f} MB",
                    "분류": category
                })
            
            df = pd.DataFrame(file_data)
            st.table(df)
        else:
            st.info("📂 저장된 파일이 없습니다.")
    else:
        st.info("📂 문서 저장 디렉토리가 없습니다.")

    # 수동 리로드 버튼
    if st.button("🔄 학습 문서 정보 리로드", type="secondary"):
        reload_backend_indexes()

def show_chat_page():
    """채팅 페이지를 표시합니다."""
    # 내부 로고 영역
    st.markdown("""
        <div style="background-color:#ffffff; padding: 1rem 2rem; border-radius: 6px; margin-bottom: 1rem; display: flex; align-items: center; border: 1px solid #e5e7eb; box-shadow: 0 2px 4px rgba(0,0,0,0.04);">
            <img src="https://img.icons8.com/ios-filled/50/2d9bf0/document--v1.png" width="24px" style="margin-right: 10px;" />
            <span style="font-size: 1.3rem; font-weight: bold; color: #111827;">DocInsight AI</span>
        </div>
        <div style="color: #666666; font-size: 0.95rem; margin-bottom: 1.5rem;">
           다양한 문서를 학습한 AI가 질문에 답변합니다.
        </div>
    """, unsafe_allow_html=True)

    display_chat_history()

    user_input = st.chat_input("서울시 도시가스 요금 산정 방식은?")
    if user_input:
        send_message(user_input)
        st.rerun()

def main():
    initialize_session_state()

    with st.sidebar:
        st.markdown("""
            <style>
            .sidebar-title {
                font-size: 1.3rem;
                font-weight: bold;
                margin-bottom: 1rem;
                color: white !important;
            }
            </style>
            <div class="sidebar-title">문서관리</div>
        """, unsafe_allow_html=True)
        
        # API 처리 중인지 확인
        is_processing = st.session_state.get("api_processing", False)
        
        # 문서업로드 버튼
        upload_clicked = st.button(
            "📂 문서 업로드", 
            use_container_width=True,
            disabled=is_processing,
            key="upload_button"
        )
        
        if upload_clicked and not is_processing:
            st.session_state.current_page = "upload"
            st.rerun()
        
        # 뒤로가기 버튼은 문서 업로드 화면에서만 표시
        if st.session_state.current_page == "upload":
            # 문서 업로드 버튼 바로 아래에 뒤로가기 버튼 배치
            back_clicked = st.button(
                "← 뒤로가기", 
                use_container_width=True,
                disabled=is_processing,
                key="back_button"
            )
            
            if back_clicked and not is_processing:
                st.session_state.current_page = "chat"
                st.rerun()

    # 상단 전체 헤더 영역 어둡게 처리
    st.markdown("""
        <style>
        header[data-testid="stHeader"] {
            background-color: #111827;
        }

        section[data-testid="stSidebar"] {
            width: 187px !important;
            background-color: #111827 !important;
            color: white !important;
            overflow-y: hidden !important;  /* 사이드바 세로 스크롤 비활성화 */
            overflow-x: hidden !important;  /* 사이드바 가로 스크롤 비활성화 */
            height: 100vh !important;       /* 사이드바 높이를 화면 높이로 고정 */
            max-height: 100vh !important;   /* 최대 높이 제한 */
        }

        /* 사이드바 내부 컨테이너도 스크롤 방지 */
        section[data-testid="stSidebar"] > div {
            overflow-y: hidden !important;
            overflow-x: hidden !important;
            height: 100vh !important;
            max-height: 100vh !important;
        }

        /* 사이드바 내부 모든 요소들의 스크롤 방지 */
        section[data-testid="stSidebar"] * {
            overflow-y: visible !important;
            overflow-x: visible !important;
        }

        /* 사이드바 버튼이 disabled 상태에서도 색상 유지 */
        section[data-testid="stSidebar"] .stButton > button {
            background-color: #333333 !important;
            color: white !important;
            border-radius: 8px !important;
            padding: 0.5rem 1rem !important;
        }
        section[data-testid="stSidebar"] .stButton > button:hover {
            background-color: #555555 !important;
        }
        section[data-testid="stSidebar"] .stButton > button:disabled {
            background-color: #333333 !important;
            color: white !important;
            opacity: 0.6 !important;
            cursor: not-allowed !important;
        }

        /* 사이드바 토글 버튼 강조 */
        button[data-testid="collapsedControl"] {
            border: 2px solid #9ca3af !important;
            border-radius: 50% !important;
            background-color: #f3f4f6 !important;
            width: 36px !important;
            height: 36px !important;
            display: flex;
            align-items: center;
            justify-content: center;
            position: fixed;
            top: 1.25rem;
            left: 0.75rem;
            z-index: 10000;
        }
        button[data-testid="collapsedControl"] svg {
            stroke: #374151 !important;
            width: 20px !important;
            height: 20px !important;
        }
        </style>
    """, unsafe_allow_html=True)

    # 현재 페이지에 따라 다른 내용 표시
    if st.session_state.current_page == "upload":
        show_upload_page()
    else:
        show_chat_page()

    st.markdown("""
        <style>
        html, body, [class*="css"]  {
            font-family: 'Pretendard', sans-serif !important;
        }
        /* 메인 콘텐츠 영역의 버튼만 스타일 적용 (사이드바 제외) */
        .main .stButton button {
            background-color: #333333 !important;
            color: white !important;
            border-radius: 8px !important;
            padding: 0.5rem 1rem !important;
        }
        .main .stButton button:hover {
            background-color: #555555 !important;
        }
        </style>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
