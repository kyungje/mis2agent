import streamlit as st
import requests
import json
from typing import List, Dict, Any
import time
import re
import os
import tempfile
import shutil
from pathlib import Path

# build_faiss_with_metadata.py에서 필요한 모듈들 import
import pdfplumber
from docx import Document
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
import gc
from dotenv import load_dotenv

# .env 파일 로드
load_dotenv()

# API 엔드포인트 설정
API_URL = "http://localhost:8000/chat"
RELOAD_API_URL = "http://localhost:8000/reload-indexes"

def latex_to_text(text):
    """
    LaTeX 수식을 사람이 읽는 텍스트 수식으로 변환
    예: \frac{a}{b} → (a) / (b)
    """
    # \frac 변환 함수
    def frac_repl(match):
        return f"({match.group(1)}) / ({match.group(2)})"

    # LaTeX 블록(\[...\], $$...$$, $...$)을 찾아서 변환
    def latex_block_repl(match):
        latex_expr = match.group(1)
        # \left와 \right 제거 (괄호 크기 조정 명령어) - 더 포괄적으로 처리
        latex_expr = re.sub(r'\\left\s*\(', '(', latex_expr)
        latex_expr = re.sub(r'\\right\s*\)', ')', latex_expr)
        latex_expr = re.sub(r'\\left\s*\[', '[', latex_expr)
        latex_expr = re.sub(r'\\right\s*\]', ']', latex_expr)
        latex_expr = re.sub(r'\\left\s*\\{', '{', latex_expr)
        latex_expr = re.sub(r'\\right\s*\\}', '}', latex_expr)
        # \frac 변환
        latex_expr = re.sub(r'\\frac\{(.+?)\}\{(.+?)\}', frac_repl, latex_expr)
        # \times 변환
        latex_expr = latex_expr.replace(r'\times', '×')
        # 중괄호 제거
        latex_expr = latex_expr.replace('{', '').replace('}', '')
        return latex_expr

    # \[ ... \] 블록 변환
    text = re.sub(r'\\\[(.*?)\\\]', lambda m: latex_block_repl(m), text, flags=re.DOTALL)
    # $$ ... $$ 블록 변환
    text = re.sub(r'\$\$(.*?)\$\$', lambda m: latex_block_repl(m), text, flags=re.DOTALL)
    # $ ... $ 블록 변환
    text = re.sub(r'\$(.*?)\$', lambda m: latex_block_repl(m), text, flags=re.DOTALL)

    # 인라인 변환 (혹시 남아있을 경우)
    # \left와 \right 제거 - 더 포괄적으로 처리
    text = re.sub(r'\\left\s*\(', '(', text)
    text = re.sub(r'\\right\s*\)', ')', text)
    text = re.sub(r'\\left\s*\[', '[', text)
    text = re.sub(r'\\right\s*\]', ']', text)
    text = re.sub(r'\\left\s*\\{', '{', text)
    text = re.sub(r'\\right\s*\\}', '}', text)
    # \frac 변환
    text = re.sub(r'\\frac\{(.+?)\}\{(.+?)\}', frac_repl, text)
    text = text.replace(r'\times', '×')
    text = text.replace('{', '').replace('}', '')

    return text

# === build_faiss_with_metadata.py에서 가져온 함수들 ===

# 경로 설정
BASE_DIR = Path(__file__).parent.parent / "vectordb"
DOCS_DIR = BASE_DIR / "docs"
DB_DIR = BASE_DIR / "db"

# 인덱스 분류별 경로 설정
GAS_INDEX_DIR = DB_DIR / "gas_index"
POWER_INDEX_DIR = DB_DIR / "power_index"
OTHER_INDEX_DIR = DB_DIR / "other_index"

# 필요한 폴더 생성
DB_DIR.mkdir(exist_ok=True)
GAS_INDEX_DIR.mkdir(exist_ok=True)
POWER_INDEX_DIR.mkdir(exist_ok=True)
OTHER_INDEX_DIR.mkdir(exist_ok=True)

# OpenAI 임베딩 모델 초기화
def create_embedding_model():
    """동적으로 배치 크기를 조정하여 임베딩 모델을 생성합니다."""
    chunk_size = 1000
    
    try:
        embedding_model = OpenAIEmbeddings(
            model="text-embedding-3-small",
            chunk_size=chunk_size,
            max_retries=3
        )
        return embedding_model
    except Exception as e:
        if "max_tokens_per_request" in str(e):
            chunk_size = 500
            st.info(f"토큰 제한으로 인해 배치 크기를 {chunk_size}로 조정합니다.")
            return OpenAIEmbeddings(
                model="text-embedding-3-small",
                chunk_size=chunk_size,
                max_retries=3
            )
        else:
            raise e

# 텍스트 분할기 초기화
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=400,
    chunk_overlap=80,
    length_function=len,
    separators=["\n\n", "\n", " ", ""]
)

# 문서 읽기 함수들
def read_docx(path):
    """DOCX 문서를 읽어서 텍스트를 추출합니다."""
    doc = Document(path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    text = latex_to_text(text)
    return text

def read_pdf(path):
    """PDF 문서를 읽어서 전체 텍스트를 추출합니다."""
    full_text = ""
    with pdfplumber.open(path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            try:
                page_text = page.extract_text()
                if page_text:
                    page_text = latex_to_text(page_text)
                    page_text = normalize_text(page_text)
                    full_text += page_text + "\n"
            except Exception as e:
                st.warning(f"페이지 {page_num} 처리 중 오류: {str(e)}")
    return full_text

def read_txt(path):
    """TXT 파일을 읽어서 텍스트를 추출합니다."""
    with open(path, 'r', encoding='utf-8') as f:
        text = f.read()
    text = latex_to_text(text)
    return text

# 텍스트 정규화
def normalize_text(text):
    """텍스트 정규화 (한글-영문 공백, 수식 기호 보존)"""
    # 한글-영문 공백 정리
    text = re.sub(r'([가-힣])([A-Za-z0-9])', r'\1 \2', text)
    text = re.sub(r'([A-Za-z0-9])([가-힣])', r'\1 \2', text)

    # 수식 기호 보존: √ ± ≈ ∞ × ÷ π ² ³ ^ / = % 등
    math_symbols = "√±≈∞×÷π²³^=/%"

    # 더 관대한 정규화: 한글, 영문, 숫자, 공백, 문장부호, 수식 기호만 유지
    # 한글: 가-힣
    # 영문: A-Za-z
    # 숫자: 0-9
    # 공백: \s
    # 문장부호: .,!?;:-()[]{}
    # 수식 기호: √±≈∞×÷π²³^=/%
    # 추가 허용 문자: + (플러스 기호)
    
    allowed_pattern = r'[가-힣A-Za-z0-9\s.,!?;:\-\(\)\[\]\{\}' + re.escape(math_symbols + '+') + r']'
    text = re.sub(rf'[^{allowed_pattern}]', ' ', text)

    # 연속 공백 정리
    text = re.sub(r'\s+', ' ', text)

    return text.strip()

# 파일명 정규화 함수
def normalize_filename(file_name):
    """파일명을 정규화하여 분류에 사용합니다."""
    # 파일 확장자 제거
    name_without_ext = os.path.splitext(file_name)[0]
    
    # 특수문자 제거 (하이픈, 언더스코어, 플러스 등은 공백으로 변환)
    normalized = re.sub(r'[\[\]\(\)\+\-\_]', ' ', name_without_ext)
    
    # 연속 공백을 단일 공백으로 변환
    normalized = re.sub(r'\s+', ' ', normalized)
    
    # 앞뒤 공백 제거
    normalized = normalized.strip()
    
    return normalized

# 파일 분류 함수
def classify_file(file_name):
    """파일명에 따라 분류를 결정합니다."""
    st.info(f"🔍 파일 분류 중: '{file_name}'")
    
    # 파일명 정규화
    normalized_name = normalize_filename(file_name)
    st.info(f"  📝 정규화된 파일명: '{normalized_name}'")
    
    if '도시가스' in normalized_name:
        st.info(f"  ✅ '도시가스' 키워드 발견 → Gas 분류")
        return 'gas'
    elif '전력' in normalized_name:
        st.info(f"  ✅ '전력' 키워드 발견 → Power 분류")
        return 'power'
    else:
        st.info(f"  ⚠️ 키워드 없음 → Other 분류")
        return 'other'

def get_index_dir(category):
    """분류에 따른 인덱스 디렉토리를 반환합니다."""
    if category == 'gas':
        return GAS_INDEX_DIR
    elif category == 'power':
        return POWER_INDEX_DIR
    else:
        return OTHER_INDEX_DIR

def extract_metadata_from_filename(filename):
    base_name = os.path.splitext(filename)[0]
    parts = re.split(r'[_\-\s]', base_name)

    version = next((p for p in parts if re.match(r'20\d{2}', p)), None)
    region_list = ['서울', '부산', '대구', '광주', '인천', '대전', '울산','경기도', '강원도', '충청북도', '충청남도' ,'전라남도','전북특별자치도', '경상남도', '경상북도']
    region = next((p for p in parts if p in region_list), None)
    organization_map = {'도시가스': '도시가스', '전력': '전력'}
    organization = next((p for p in organization_map if p in parts), "기타")

    return {
        "version": version,
        "region": region,
        "organization": organization,
        "title": base_name
    }

# 문서 처리 및 청킹
def process_document(file_path):
    """문서를 읽고 LangChain 텍스트 분할기를 사용하여 청킹합니다."""
    filename = os.path.basename(file_path)
    
    # 문서 읽기
    if file_path.endswith(".pdf"):
        text = read_pdf(file_path)
    elif file_path.endswith(".docx"):
        text = read_docx(file_path)
    elif file_path.endswith(".txt"):
        text = read_txt(file_path)
    else:
        return []
    
    # 텍스트가 비어있으면 처리 중단
    if not text:
        st.warning(f"⚠️ {filename}에서 텍스트를 추출하지 못했습니다. 건너뜁니다.")
        return []
    
    st.info(f"  📄 원본 텍스트 길이: {len(text)}자")
    
    # 텍스트 정규화
    original_text = text
    text = normalize_text(text)
    st.info(f"  📄 정규화 후 텍스트 길이: {len(text)}자")
    
    # 정규화 과정에서 텍스트가 너무 많이 제거되었는지 확인
    if len(text) < len(original_text) * 0.1:  # 90% 이상 제거된 경우
        st.warning(f"  ⚠️ 텍스트가 너무 많이 제거되었습니다. 원본: {len(original_text)}자 → 정규화: {len(text)}자")
        # 정규화를 건너뛰고 원본 텍스트 사용
        text = original_text
        st.info(f"  🔄 원본 텍스트를 사용합니다.")
    
    # LangChain 텍스트 분할기 사용
    chunks = text_splitter.split_text(text)
    st.info(f"  📄 청킹 후 청크 수: {len(chunks)}개")
    
    # 파일명에서 추가 메타데이터 추출
    additional_metadata = extract_metadata_from_filename(filename)
    
    # 메타데이터가 포함된 LangChain 문서로 변환
    documents = []
    for i, chunk in enumerate(chunks):
        if len(chunk.strip()) > 50:  # 최소 길이 필터 (50자 이상)
            # 기본 메타데이터와 추가 메타데이터를 결합
            metadata = {
                "source": filename,
                "chunk_id": i,
                "file_path": file_path,
                "version": additional_metadata["version"],
                "region": additional_metadata["region"],
                "organization": additional_metadata["organization"],
                "title": additional_metadata["title"]
            }
            doc = LCDocument(
                page_content=chunk,
                metadata=metadata
            )
            documents.append(doc)
        else:
            st.info(f"  ⚠️ 청크 {i}가 너무 짧아 제외됨: {len(chunk.strip())}자")
    
    st.info(f"[CHUNKS] {filename} → {len(documents)}개 생성")
    return documents

# 벡터 인덱스 구축
def build_vector_index_from_uploaded_files(uploaded_files):
    """업로드된 파일들로부터 벡터 인덱스를 구축합니다."""
    if not uploaded_files:
        st.warning("업로드된 파일이 없습니다.")
        return False
    
    st.info("📂 문서 인덱싱 시작")
    
    # 문서 저장 디렉토리 생성
    docs_dir = Path("/Users/a07198/IdeaProjects/MIS2/src/vectordb/docs")
    docs_dir.mkdir(parents=True, exist_ok=True)
    st.info(f"📁 문서 저장 디렉토리: {docs_dir}")
    
    # 임베딩 모델 초기화
    embedding_model = create_embedding_model()
    
    # 분류별로 문서 그룹화
    gas_documents = []
    power_documents = []
    other_documents = []
    
    total_files = len(uploaded_files)
    
    # 파일별로 처리 및 분류
    for i, uploaded_file in enumerate(uploaded_files, 1):
        st.info(f"[{i}/{total_files}] 처리 중: {uploaded_file.name}")
        
        # 실제 파일을 docs 디렉토리에 저장
        file_path = docs_dir / uploaded_file.name
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getvalue())
        
        st.info(f"  💾 파일 저장: {file_path}")
        
        try:
            # 파일 분류
            category = classify_file(uploaded_file.name)
            st.info(f"  → 분류: {category}")
            
            documents = process_document(str(file_path))
            
            # 분류별로 문서 추가
            if category == 'gas':
                gas_documents.extend(documents)
            elif category == 'power':
                power_documents.extend(documents)
            else:
                other_documents.extend(documents)
            
            # 메모리 최적화
            del documents
            gc.collect()
            
        except Exception as e:
            st.error(f"  ❌ 파일 처리 오류: {e}")
            # 오류 발생 시 저장된 파일 삭제
            if file_path.exists():
                file_path.unlink()
            continue
    
    st.info(f"\n📊 전체 문서 수: {len(uploaded_files)}개")
    st.info(f"🔖 Gas 문서: {len(gas_documents)}개")
    st.info(f"🔖 Power 문서: {len(power_documents)}개")
    st.info(f"🔖 Other 문서: {len(other_documents)}개")
    
    # 분류별로 인덱스 생성
    categories = [
        ('gas', gas_documents, 'Gas'),
        ('power', power_documents, 'Power'),
        ('other', other_documents, 'Other')
    ]
    
    success_count = 0
    for category, documents, category_name in categories:
        if len(documents) == 0:
            st.warning(f"⚠️ {category_name} 문서가 없어 인덱스를 생성하지 않습니다.")
            continue
            
        st.info(f"\n🔧 {category_name} 인덱스 생성 중... (문서 수: {len(documents)}개)")
        
        # 배치 단위로 임베딩 처리
        try:
            vectorstore = FAISS.from_documents(documents, embedding_model)
        except Exception as e:
            if "max_tokens_per_request" in str(e):
                st.info(f"토큰 제한으로 인해 배치 크기를 500으로 조정합니다...")
                try:
                    medium_embedding_model = OpenAIEmbeddings(
                        model="text-embedding-3-small",
                        chunk_size=500,
                        max_retries=3
                    )
                    vectorstore = FAISS.from_documents(documents, medium_embedding_model)
                except Exception as e2:
                    if "max_tokens_per_request" in str(e2):
                        st.info(f"토큰 제한으로 인해 배치 크기를 100으로 조정합니다...")
                        small_embedding_model = OpenAIEmbeddings(
                            model="text-embedding-3-small",
                            chunk_size=100,
                            max_retries=3
                        )
                        vectorstore = FAISS.from_documents(documents, small_embedding_model)
                    else:
                        raise e2
            else:
                raise e

        # 인덱스 저장
        index_dir = get_index_dir(category)
        try:
            vectorstore.save_local(str(index_dir))
            st.success(f"💾 {category_name} 인덱스 저장 완료: {index_dir}")
            success_count += 1
        except Exception as e:
            st.error(f"❌ {category_name} 인덱스 저장 오류: {e}")
    
    if success_count > 0:
        st.success(f"\n🎉 {success_count}개 분류별 인덱스 생성 완료")
        return True
    else:
        st.error("❌ 인덱스 생성에 실패했습니다.")
        return False

def reload_backend_indexes():
    """백엔드의 인덱스를 다시 로드합니다."""
    try:
        st.info("🔄 백엔드 인덱스 리로드 중...")
        response = requests.post(RELOAD_API_URL)
        response.raise_for_status()
        
        result = response.json()
        if result.get("success"):
            st.success("✅ 백엔드 인덱스 리로드 완료!")
            return True
        else:
            st.error(f"❌ 백엔드 인덱스 리로드 실패: {result.get('message', '알 수 없는 오류')}")
            return False
    except Exception as e:
        st.error(f"❌ 백엔드 인덱스 리로드 중 오류: {str(e)}")
        st.warning("⚠️ FastAPI 서버가 실행 중인지 확인해주세요.")
        return False

def initialize_session_state():
    """세션 상태 초기화"""
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "streaming" not in st.session_state:
        st.session_state.streaming = False

def display_chat_history():
    """채팅 기록 표시"""
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

def stream_response(response_text: str, loading_placeholder):
    """스트리밍 방식으로 응답 표시"""
    message_placeholder = st.empty()
    full_response = ""
    
    # 로딩 메시지 제거
    loading_placeholder.empty()
    
    # 문자를 하나씩 표시
    for char in response_text:
        full_response += char
        message_placeholder.markdown(full_response + "▌")
        time.sleep(0.01)  # 타이핑 효과를 위한 지연
    
    # 최종 응답 표시
    message_placeholder.markdown(full_response)
    return full_response

def send_message(user_input: str):
    """메시지 전송 및 응답 처리"""
    if not user_input:
        return

    # 사용자 메시지를 먼저 화면에 표시
    with st.chat_message("user"):
        st.markdown(user_input)
    
    # 로딩 메시지 표시를 위한 플레이스홀더
    loading_placeholder = st.empty()
    with loading_placeholder:
        st.markdown('🤖 AI가 응답을 생성하고 있습니다...')
    
    # 사용자 메시지를 세션에 추가
    st.session_state.messages.append({"role": "user", "content": user_input})
    
    # API 요청 데이터 준비
    request_data = {
        "messages": [
            {"role": msg["role"], "content": msg["content"]}
            for msg in st.session_state.messages
        ]
    }
    
    try:
        # API 호출
        response = requests.post(API_URL, json=request_data)
        response.raise_for_status()
        
        # 응답 처리
        assistant_response = response.json()["response"]
        
        # === 여기에서 수식 변환 적용 ===
        assistant_response = latex_to_text(assistant_response)
        
        # 스트리밍 방식으로 응답 표시
        with st.chat_message("assistant"):
            streamed_response = stream_response(assistant_response, loading_placeholder)
            st.session_state.messages.append({"role": "assistant", "content": streamed_response})
        
    except Exception as e:
        loading_placeholder.empty()
        st.error(f"Error: {str(e)}")

def main():
    st.set_page_config(page_title="AI Agent Chat", page_icon="🤖", layout="wide")
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "streaming" not in st.session_state:
        st.session_state.streaming = False

    tab1, tab2 = st.tabs(["💬 채팅", "📁 문서 업로드"])

    # 채팅 탭
    with tab1:
        with st.sidebar:
            st.header("About")
            st.markdown("""
            This is a chat interface built with Streamlit.
            It uses OpenAI's GPT model through a FastAPI backend.
            """)
            if st.button("Clear Chat"):
                st.session_state.messages = []
                st.rerun()

        # 항상 채팅 입력창이 하단에 고정
        display_chat_history()
        user_input = st.chat_input("메시지를 입력하세요", key="chat_input", disabled=False)
        if user_input:
            send_message(user_input)
            st.rerun()

    # 문서 업로드 탭 (복원)
    with tab2:
        st.header("📁 문서 업로드 및 인덱스 생성")
        st.markdown("""
        ### 지원 파일 형식
        - **PDF** (.pdf)
        - **Word 문서** (.docx) 
        - **텍스트 파일** (.txt)

        ### 파일 분류 규칙
        - 파일명에 **'도시가스'** 포함 → Gas 분류
        - 파일명에 **'전력'** 포함 → Power 분류  
        - 기타 → Other 분류
        """)
        uploaded_files = st.file_uploader(
            "문서 파일을 선택하세요",
            type=['pdf', 'docx', 'txt'],
            accept_multiple_files=True,
            help="여러 파일을 동시에 업로드할 수 있습니다."
        )
        if uploaded_files:
            st.info(f"📁 {len(uploaded_files)}개 파일이 선택되었습니다:")
            for file in uploaded_files:
                category = classify_file(file.name)
                st.write(f"- **{file.name}** → **{category}** 분류")
                if category == 'other':
                    st.warning(f"  ⚠️ '{file.name}'이 'other'로 분류되었습니다. 파일명에 '도시가스' 또는 '전력'이 포함되어 있는지 확인해주세요.")
        if st.button("🚀 인덱스 생성 시작", type="primary", disabled=not uploaded_files):
            try:
                with st.spinner("문서를 처리하고 인덱스를 생성하고 있습니다..."):
                    success = build_vector_index_from_uploaded_files(uploaded_files)
                    if success:
                        st.balloons()
                        st.success("✅ 인덱스 생성이 완료되었습니다!")
                        
                        # 백엔드 인덱스 리로드
                        reload_success = reload_backend_indexes()
                        if reload_success:
                            st.success("🎉 모든 작업이 완료되었습니다! 이제 채팅 탭에서 질문할 수 있습니다.")
                        else:
                            st.warning("⚠️ 인덱스는 생성되었지만 백엔드 리로드에 실패했습니다. FastAPI 서버를 재시작하거나 수동으로 리로드해주세요.")
                    else:
                        st.error("❌ 인덱스 생성에 실패했습니다. 다시 시도해주세요.")
            except Exception as e:
                st.error(f"❌ 인덱스 생성 중 오류가 발생했습니다: {str(e)}")
                st.error("API 서버가 실행되지 않았을 수 있습니다. FastAPI 서버를 먼저 실행해주세요.")
        st.header("📊 기존 인덱스 정보")
        col1, col2, col3 = st.columns(3)
        with col1:
            if GAS_INDEX_DIR.exists() and any(GAS_INDEX_DIR.iterdir()):
                st.success("✅ Gas 인덱스 존재")
            else:
                st.warning("⚠️ Gas 인덱스 없음")
        with col2:
            if POWER_INDEX_DIR.exists() and any(POWER_INDEX_DIR.iterdir()):
                st.success("✅ Power 인덱스 존재")
            else:
                st.warning("⚠️ Power 인덱스 없음")
        with col3:
            if OTHER_INDEX_DIR.exists() and any(OTHER_INDEX_DIR.iterdir()):
                st.success("✅ Other 인덱스 존재")
            else:
                st.warning("⚠️ Other 인덱스 없음")
        
        # 수동 리로드 버튼
        st.header("🔄 백엔드 인덱스 관리")
        if st.button("🔄 백엔드 인덱스 리로드", type="secondary"):
            reload_backend_indexes()
        
        # 저장된 문서 파일 목록 표시
        st.header("📁 저장된 문서 파일")
        docs_dir = Path("/Users/a07198/IdeaProjects/MIS2/src/vectordb/docs")
        if docs_dir.exists() and any(docs_dir.iterdir()):
            files = list(docs_dir.glob("*"))
            if files:
                st.info(f"📂 총 {len(files)}개 파일이 저장되어 있습니다:")
                for file in sorted(files):
                    file_size = file.stat().st_size
                    file_size_mb = file_size / (1024 * 1024)
                    category = classify_file(file.name)
                    st.write(f"- **{file.name}** ({category} 분류, {file_size_mb:.2f} MB)")
            else:
                st.info("📂 저장된 파일이 없습니다.")
        else:
            st.info("📂 문서 저장 디렉토리가 없습니다.")

if __name__ == "__main__":
    main() 